import docx
from django.core.management.base import BaseCommand
from django.contrib.auth import authenticate, login
from django.contrib import messages
from .models import Vendedor, Zona, Transporte, Impostos, Pagamento, Modalidade, Precos, Cliente1, Artigo, \
    DocumentoContador, Recibo, DocumentoTemp, Empresa, Moeda, TaxReason, \
    TempArtigos, DocumentoFinalizado, FinArtigos, ReciboLinhas, DocumentoFinalizadoContador
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.decorators import login_required
from functools import wraps

@login_required
def webapp_view(request):
    # REGRA: Se tem email e não verificou o código, manda para o check_mfa
    if request.user.email and not request.user.is_verified():
        return redirect('check_mfa_status')

    empresa = Empresa.objects.filter(user=request.user).first()
    if not empresa:
        return redirect('completar_empresa')

    context = {
        'user': request.user,
        'empresa': empresa.nome,
    }
    return render(request, 'webapp.html', context)

@login_required
def check_mfa_status(request):
    user = request.user

    if not user.email:
        return redirect('webapp_home')

    if not user.is_verified():
        return redirect('otp_verify_view')

    return redirect('webapp_home')


from django.contrib.auth.decorators import login_required

from django_otp import login as otp_login


@login_required
def otp_verify_view(request):
    from django_otp.plugins.otp_email.models import EmailDevice
    from django_otp.util import random_hex

    # 1. Garante que o utilizador tem um "dispositivo" de e-mail registado
    device, created = EmailDevice.objects.get_or_create(
        user=request.user,
        name="default",
        defaults={'email': request.user.email}
    )

    if not request.user.email:
        messages.error(request, "O teu utilizador não tem e-mail configurado.")
        return redirect('webapp_home')

    # 2. Se o utilizador já validou o OTP nesta sessão, entra direto
    if request.user.is_verified():
        return redirect('webapp_home')

    if request.method == "POST":
        token = request.POST.get("otp_token")
        if device.verify_token(token):
            from django_otp import login as otp_login
            otp_login(request, device)
            messages.success(request, "Bem-vindo! Verificação concluída.")
            return redirect('webapp_home')
        else:
            messages.error(request, "Código inválido. Tente novamente.")
    else:
        # 3. GERAR E ENVIAR O CÓDIGO REAL
        try:
            # O django-otp-email tem um método próprio para gerar e enviar
            device.generate_challenge()

            messages.info(request, f"Enviámos um código de 6 dígitos para {request.user.email}")
            print(f"OTP enviado com sucesso para {request.user.email}")
        except Exception as e:
            messages.error(request, "Erro ao enviar o código. Tenta mais tarde.")
            print(f"Erro ao enviar OTP: {e}")

    return render(request, 'account/verify.html')

from .forms import EmpresaForm

@login_required
def completar_registo_empresa(request):
    if hasattr(request.user, 'empresa'):
        return redirect('webapp_home')

    if request.method == 'POST':
        form = EmpresaForm(request.POST)
        if form.is_valid():
            empresa = form.save(commit=False)
            empresa.user = request.user
            empresa.save()
            return redirect('webapp_home')
    else:
        form = EmpresaForm()

    return render(request, 'account/completar_empresa.html', {'form': form})

from django.contrib.auth import logout
@login_required
def logout_view(request):
    logout(request)
    return redirect('/accounts/login/')

from .decorators import empresa_obrigatoria
@login_required
@empresa_obrigatoria
def clientes_json(request):
    clientes = Cliente1.objects.filter(empresa=request.empresa).values(
        "id_cliente", "codigo", "nome", "contribuinte", "morada1", "codigo_postal", "concelho"
    )

    return JsonResponse({'data': list(clientes)})
@login_required
@empresa_obrigatoria
def faturas_json(request):
    TIPO_EXTENSO = {
        'FT': 'Fatura',
        'FR': 'Fatura-Recibo',
        'FS': 'Fatura Simplificada',
        'NC': 'Nota de Crédito',
    }

    documentos = (
        DocumentoFinalizado.objects
        .filter(empresa=request.empresa, tipo__in=TIPO_EXTENSO.keys())
        .order_by('-id', '-data_emissao')
    )

    documentos_temp = (
        DocumentoTemp.objects
        .filter(empresa=request.empresa, tipo__in=TIPO_EXTENSO.keys())
        .select_related('cliente')
        .order_by('-id', '-criado_em')
    )

    data = []
    for doc in documentos_temp:
        cliente = doc.cliente
        valor_total = float(doc.valor_total) if getattr(doc, 'valor_total', None) else 0

        numero_completo = f"{doc.tipo}/{doc.serie}/{doc.ano}/{doc.numero}" if getattr(doc, 'serie',
                                                                                          None) else f"Temp-{doc.id}"

        data.append({
            "id_documento": doc.id,
            "tipo": f"{TIPO_EXTENSO.get(doc.tipo, doc.tipo)} (Temp)",
            "numero_doc": doc.numero,
            "numero": numero_completo,
            "cliente_id": cliente.id_cliente if cliente else None,
            "cliente_nome": cliente.nome if cliente else "",
            "data_emissao": doc.data_emissao.strftime('%Y-%m-%d') if doc.data_emissao else "",
            "vencimento": doc.data_vencimento if doc.data_vencimento else "",
            "valor_total": valor_total,
            "total_pago": 0,
            "restante": valor_total,
            "estado_pagamento":'Pendente',
            "estado": 'Rascunho',
            "temporario": True
        })
    for doc in documentos:
        numero_completo = f"{doc.tipo}/{doc.serie}/{doc.ano}/{doc.numero}"

        data.append({
            "id_documento": doc.id,
            "tipo": TIPO_EXTENSO.get(doc.tipo, doc.tipo),
            "numero_doc": doc.numero,
            "numero": numero_completo,
            "cliente_id": doc.cliente_id,
            "cliente_nome": doc.cliente_nome,
            "data_emissao": doc.data_emissao.strftime('%Y-%m-%d') if doc.data_emissao else "",
            "vencimento": doc.data_vencimento,
            "valor_total": float(doc.valor_total),
            "total_pago": float(doc.total_pago),
            "restante": float(doc.valor_total - doc.total_pago),
            "estado_pagamento": doc.estado_pagamento,
            "estado": doc.estado,
        })


    return JsonResponse(data, safe=False)


from django.db.models import Max
from django.db.models.functions import Cast, Substr
from django.db.models import IntegerField
@login_required
@empresa_obrigatoria
def proximo_codigo_cliente(request):
    ultimo = Cliente1.objects.filter(empresa=request.empresa).annotate(
        codigo_num=Cast(Substr('codigo', 2), IntegerField())
    ).aggregate(max_codigo=Max('codigo_num'))['max_codigo']

    if ultimo:
        numero = ultimo + 1
    else:
        numero = 1

    proximo = f"C{numero:03d}"
    return JsonResponse({"codigo": proximo})

@login_required
@empresa_obrigatoria
def proximo_codigo_artigo(request):
    # Filtra os artigos APENAS da empresa do utilizador logado
    ultimo = Artigo.objects.filter(empresa=request.empresa).aggregate(
        max_codigo=Max('id_artigo')
    )['max_codigo']

    numero = (ultimo or 0) + 1

    return JsonResponse({"codigo": numero})


def validar_nif_portugal(nif):
    nif = str(nif)
    if not nif.isdigit() or len(nif) != 9:
        return False

    if nif[0] not in '123456789':
        return False

    # Prefixos específicos de 2 dígitos
    if nif[0] == '4' and nif[1] != '5': return False

    soma = 0
    for i in range(8):
        soma += int(nif[i]) * (9 - i)

    resto = soma % 11
    digito_controlo = 0 if resto < 2 else 11 - resto
    return int(nif[8]) == digito_controlo

from django.core.validators import validate_email
from django.core.exceptions import ObjectDoesNotExist
@login_required
@empresa_obrigatoria
def validar_dados_cliente(request, post_data, cliente_id=None, is_edit=False):
    """
    Valida os dados do cliente para garantir conformidade com a AT e integridade da BD.
    """
    regras = {
        'nome': ('Nome', 200),
        'contribuinte': ('NIF', 12),
        'morada1': ('Morada 1', 100),
        'morada2': ('Morada 2', 100),
        'codigo_postal': ('Código Postal', 8),
        'distrito': ('Distrito', 35),
        'concelho': ('Concelho', 35),
        'telemovel': ('Telemóvel', 12),
        'pais': ('Country',40),
        'sigla': ('Sigla', 2),
        'email': ('Email', 200),
    }

    obrigatorios = ['nome', 'morada1', 'codigo_postal', 'pais', 'distrito', 'concelho', 'vendedor', 'impostos']

    if is_edit:
        obrigatorios = [campo for campo in obrigatorios if campo != 'nome']

    for campo, (label, max_len) in regras.items():
        valor = post_data.get(campo, '').strip()

        if campo in obrigatorios and not valor:
            return False, f"O campo [{label}] é obrigatório."

        if len(valor) > max_len:
            return False, f"O campo [{label}] excede o limite de {max_len} caracteres."

    email = post_data.get('email', '').strip()
    if email:
        try:
            validate_email(email)
        except ValidationError:
            return False, "O formato do email introduzido é inválido."

    nif = post_data.get('contribuinte', '').replace(' ', '')
    pais = post_data.get('pais', 'PT').upper()


    if pais == "PT":
        if not nif:
            nif = "999999990"
            post_data['contribuinte'] = nif

        if nif != "999999990":
            if len(nif) != 9:
                return False, "O NIF para clientes em Portugal deve ter exatamente 9 dígitos."
            if not validar_nif_portugal(nif):
                return False, "O NIF introduzido é inválido para Portugal."
    else:
        if not nif:
            return False, "O NIF é obrigatório para clientes estrangeiros."

        query_duplicado = Cliente1.objects.filter(contribuinte=nif, empresa=request.empresa)
        if cliente_id:
            query_duplicado = query_duplicado.exclude(id_cliente=cliente_id)

        if query_duplicado.exists():
            return False, f"Já existe um cliente registado com o NIF {nif}."

    modelos_fk = [
        (Vendedor, 'vendedor', 'id_vendedor'),
        (Zona, 'zona', 'id_zona'),
        (Impostos, 'impostos', 'id_impostos'),
        (Pagamento, 'pagamento', 'id_pagamento'),
        (Modalidade, 'modalidade', 'id_modalidade'),
        (Precos, 'precos', 'id_precos'),
    ]

    for model, field_name, id_name in modelos_fk:
        val = post_data.get(field_name)
        if field_name in obrigatorios or (val and val != ""):

            # Cria o dicionário de filtros dinamicamente
            query_params = {id_name: val}

            # Apenas adiciona o filtro de empresa se o modelo realmente o tiver
            if hasattr(model, 'empresa'):
                query_params['empresa'] = request.empresa

            try:
                model.objects.get(**query_params)
            except (ObjectDoesNotExist, ValueError):
                return False, f"A opção selecionada para [{field_name.capitalize()}] é inválida."

    return True, "Validado com sucesso"

from django.urls import reverse
@login_required
@empresa_obrigatoria
def adicionar_cliente(request):
    context = {
        "vendedores": Vendedor.objects.all(),
        "zonas": Zona.objects.all(),
        "transportes": Transporte.objects.filter(empresa=request.empresa),
        "impostos_list": Impostos.objects.all(),
        "pagamentos": Pagamento.objects.all(),
        "modalidades": Modalidade.objects.all(),
        "precos_list": Precos.objects.all(),
    }

    if request.method == "POST":

        valido, mensagem = validar_dados_cliente(request, request.POST, cliente_id=None)

        if not valido:
            messages.error(request, mensagem)
            context['dados'] = request.POST
            return JsonResponse({'status': 'error', 'message': mensagem}, status=400)
        try:
            with transaction.atomic():
                ultimo = Cliente1.objects.filter(empresa=request.empresa).aggregate(max_codigo=Max('codigo'))['max_codigo']
                if ultimo:
                    numero = int(ultimo.replace("C", "")) + 1
                else:
                    numero = 1
                novo_codigo = f"C{numero:03d}"

                # ATRIBUIÇÃO À VARIÁVEL 'novo_cliente'
                novo_cliente = Cliente1.objects.create(
                    empresa=request.empresa,
                    codigo=novo_codigo,
                    nome=request.POST.get('nome').strip(),
                    contribuinte=request.POST.get('contribuinte') or "999999990",
                    morada1=request.POST.get('morada1'),
                    morada2=request.POST.get('morada2'),
                    codigo_postal=request.POST.get('codigo_postal'),
                    telemovel=request.POST.get('telemovel'),
                    email=request.POST.get('email', '').lower().strip(),
                    pais=request.POST.get('pais'),
                    sigla=request.POST.get('sigla'),
                    distrito=request.POST.get('distrito'),
                    concelho=request.POST.get('concelho'),
                    vendedor_id=request.POST.get('vendedor'),
                    zona_id=request.POST.get('zona'),
                    transporte_id=request.POST.get('transporte'),
                    impostos_id=request.POST.get('impostos'),
                    pagamento_id=request.POST.get('pagamento'),
                    modalidade_id=request.POST.get('modalidade'),
                    precos_id=request.POST.get('precos'),
                )

            url_detalhes = reverse('cliente_detalhes', kwargs={'id_cliente': novo_cliente.id_cliente})

            # Retorna a URL para o JS redirecionar
            return JsonResponse({
                'status': 'success',
                'message': 'Cliente adicionado com sucesso!',
                'redirect_url': url_detalhes
            })

        except Exception as e:
            print(f"Erro na gravação: {e}")
            messages.error(request, "Ocorreu um erro inesperado ao gravar na base de dados.")
            context['dados'] = request.POST
            return JsonResponse({'status': 'error', 'message': 'Erro ao gravar na base de dados.'}, status=500)
    return render(request, 'subsubconteudo/criar.html',context)

@login_required
@empresa_obrigatoria
def clientes_view(request):
    vendedores = Vendedor.objects.all().order_by('nome')
    clientes = Cliente1.objects.filter(empresa=request.empresa),
    return render(request, "webapp.html", {
        "vendedores": vendedores,
        "abrir_modal_cliente": False,
        "clientes": clientes
    })
@login_required
@empresa_obrigatoria
def artigos_view(request):
    artigos = Artigo.objects.filter(empresa=request.empresa),
    return render(request, "webapp.html", {
        "abrir_modal_artigo": False,
        "artigos": artigos
    })
@login_required
@empresa_obrigatoria
def registar_json(request):
    vendedores = list(Vendedor.objects.values('id_vendedor', 'nome'))
    zona = list(Zona.objects.values('id_zona', 'zona'))
    transporte = list(Transporte.objects.filter(empresa=request.empresa).values('id_transporte', 'descricao'))
    impostos = list(Impostos.objects.values('id_impostos', 'nome'))
    pagamento = list(Pagamento.objects.values('id_pagamento', 'nome'))
    modalidade = list(Modalidade.objects.values('id_modalidade', 'nome'))
    precos = list(Precos.objects.values('id_precos', 'nome'))
    return JsonResponse({
        "vendedores": vendedores,
        "zonas": zona,
        "transportes": transporte,
        "impostos": impostos,
        "pagamentos": pagamento,
        "modalidades": modalidade,
        "precos": precos,
    })


from django.http import Http404
@login_required
@empresa_obrigatoria
def cliente_dados(request, id_cliente):
    """
    Retorna os dados de um cliente específico em JSON,
    usado para preencher o modal de edição.
    """
    try:
        cliente = Cliente1.objects.filter(empresa=request.empresa).get(id_cliente=id_cliente)
    except Cliente1.DoesNotExist:
        raise Http404("Cliente não encontrado")

    data = {
        "id_cliente": cliente.id_cliente,
        "codigo": cliente.codigo,
        "nome": cliente.nome,
        "morada1": cliente.morada1,
        "morada2": cliente.morada2,
        "codigo_postal": cliente.codigo_postal,
        "telemovel": cliente.telemovel,
        "sigla": cliente.sigla,
        "contribuinte": cliente.contribuinte,
        "pais": cliente.pais,
        "distrito": cliente.distrito,
        "concelho": cliente.concelho,
        "email": cliente.email,
        "vendedor_id": cliente.vendedor.id_vendedor if cliente.vendedor else None,
        "vendedor_nome": cliente.vendedor.nome if cliente.vendedor else None,

        "zona_id": cliente.zona.id_zona if cliente.zona else None,
        "zona_nome": cliente.zona.zona if cliente.zona else None,

        "transporte_id": cliente.transporte.id_transporte if cliente.transporte else None,
        "transporte_nome": cliente.transporte.descricao if cliente.transporte else None,

        "impostos_id": cliente.impostos.id_impostos if cliente.impostos else None,
        "impostos_nome": cliente.impostos.nome if cliente.impostos else None,

        "pagamento_id": cliente.pagamento.id_pagamento if cliente.pagamento else None,
        "pagamento_nome": cliente.pagamento.nome if cliente.pagamento else None,

        "modalidade_id": cliente.modalidade.id_modalidade if cliente.modalidade else None,
        "modalidade_nome": cliente.modalidade.nome if cliente.modalidade else None,

        "precos_id": cliente.precos.id_precos if cliente.precos else None,
        "precos_nome": cliente.precos.nome if cliente.precos else None,
    }

    return JsonResponse(data)
@login_required
@empresa_obrigatoria
def artigo_dados(request, id_artigo):
    """
    Retorna os dados de um cliente específico em JSON,
    usado para preencher o modal de edição.
    """
    try:
        artigo = Artigo.objects.filter(empresa=request.empresa).get(id_artigo=id_artigo)
    except Artigo.DoesNotExist:
        raise Http404("Artigo não encontrado")

    data = {
        "codigo": artigo.id_artigo,
        "nome": artigo.nome,
        "descricao": artigo.descricao,
        "preco": artigo.preco,
        "taxa": artigo.taxa,
    }

    return JsonResponse(data)

@login_required
@empresa_obrigatoria
def cliente_editar(request, id_cliente, ):
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Método não permitido'}, status=405)

    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'error': 'JSON inválido'})

    valido, mensagem = validar_dados_cliente(request, data, cliente_id=id_cliente, is_edit=True)

    if not valido:
        print(f"DEBUG: Validação falhou! Motivo: {mensagem}")
        return JsonResponse({'success': False, 'error': mensagem}, status=400)
    try:
        cliente = Cliente1.objects.filter(empresa=request.empresa).get(id_cliente=id_cliente)
    except Cliente1.DoesNotExist:
        return JsonResponse({'success': False, 'error': "Cliente não encontrado"}, status=404)

    cliente.morada1 = data.get('morada1', cliente.morada1)
    cliente.morada2 = data.get('morada2', cliente.morada2)
    cliente.codigo_postal = data.get('codigo_postal', cliente.codigo_postal)
    cliente.telemovel = data.get('telemovel', cliente.telemovel)
    cliente.sigla = data.get('sigla', cliente.sigla)
    cliente.contribuinte = data.get('contribuinte', cliente.contribuinte)
    cliente.pais = data.get('pais', cliente.pais)
    cliente.distrito = data.get('distrito', cliente.distrito)
    cliente.concelho = data.get('concelho', cliente.concelho)
    cliente.email = data.get('email', cliente.email)

    cliente.vendedor_id = data.get('vendedor') or None
    cliente.zona_id = data.get('zona') or None
    cliente.transporte_id = data.get('transporte') or None
    cliente.impostos_id = data.get('impostos') or None
    cliente.pagamento_id = data.get('pagamento') or None
    cliente.modalidade_id = data.get('modalidade') or None
    cliente.precos_id = data.get('precos') or None

    try:
        cliente.save()
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'Erro ao gravar na BD: {str(e)}'}, status=500)

    return JsonResponse({'success': True, 'message': 'Cliente atualizado com sucesso'})

from .models import Artigo
@login_required
@empresa_obrigatoria
def artigo_editar(request, id_artigo=None):
    if id_artigo:
        artigo = get_object_or_404(Artigo, id_artigo=id_artigo, empresa=request.empresa)
    else:
        artigo = Artigo(empresa=request.empresa)

    if request.method == 'POST':
        nome = request.POST.get('nome', '').strip()
        descricao = request.POST.get('descricao', '').strip()
        tipo = request.POST.get('tipo', '').strip()
        taxa = request.POST.get('taxa', '').strip()
        preco = request.POST.get('preco', '').strip()

        errors = []

        if not nome:
            errors.append("Nome obrigatório")
        elif len(nome) > 40:
            errors.append("Nome deve ter no máximo 40 caracteres")

        if tipo not in ['P', 'S']:
            errors.append("Tipo inválido")

        if taxa:
            try:
                t = float(taxa.replace(',', '.'))
                if t not in [0, 6, 13, 23]:
                    errors.append("Taxa IVA inválida")
            except ValueError:
                errors.append("Taxa IVA inválida")

        if preco:
            try:
                p = float(preco.replace(',', '.'))
                if p < 0:
                    errors.append("Preço deve ser positivo")
                elif round(p, 2) != p:
                    errors.append("Preço deve ter no máximo 2 casas decimais")
            except ValueError:
                errors.append("Preço inválido")

        if descricao:
            if len(descricao) > 200:
                errors.append("Descrição deve ter no máximo 200 caracteres")

        if errors:
            return JsonResponse({'success': False, 'errors': errors}, status=400)

        def clean_numeric(value):
            if not value or value.strip() == "":
                return 0.00
            return float(value.replace(',', '.'))

        artigo.nome = nome
        artigo.descricao = descricao if descricao else None
        artigo.tipo = tipo
        artigo.taxa = clean_numeric(taxa)
        artigo.preco = clean_numeric(preco)
        artigo.save()

        return JsonResponse({'success': True})

    return render(request, "subsubconteudo/criar_editar_artigo.html", {"artigo": artigo})

@login_required
@empresa_obrigatoria
@csrf_exempt
def cliente_apagar(request, id_cliente):
    if request.method != "POST":
        return JsonResponse({"success": False, "error": "Método não permitido"}, status=405)

    try:
        cliente = Cliente1.objects.filter(empresa=request.empresa).get(id_cliente=id_cliente)

        # Verifica se existem documentos associados
        existe_documento = DocumentoFinalizado.objects.filter(
            cliente_id=id_cliente,
            empresa=request.empresa
        ).exists()

        if existe_documento:
            return JsonResponse({
                "success": False,
                "has_documents": True,  # Flag explícita
                "error": "Não é possível apagar o cliente porque existem documentos associados."
            }, status=400)

        cliente.delete()
        return JsonResponse({"success": True})

    except Cliente1.DoesNotExist:
        return JsonResponse({"success": False, "error": "Cliente não encontrado"})

from django.db.models import ProtectedError

@login_required
@empresa_obrigatoria
@csrf_exempt
def artigo_apagar(request, id_artigo):
    if request.method != "POST":
        return JsonResponse({"success": False, "error": "Método não permitido"}, status=405)

    try:
        artigo = Artigo.objects.filter(empresa=request.empresa).get(id_artigo=id_artigo)
        artigo.delete()
        return JsonResponse({"success": True})

    except Artigo.DoesNotExist:
        return JsonResponse({"success": False, "error": "Artigo não encontrado."})

    except ProtectedError:
        return JsonResponse({
            "success": False,
            "error": "Não é possível eliminar este artigo porque ele já está associado a documentos ou registos financeiros."
        })

    except Exception as e:
        return JsonResponse({"success": False, "error": f"Erro inesperado: {str(e)}"})

@login_required
@empresa_obrigatoria
def artigos_json(request):
    artigos = Artigo.objects.filter(empresa=request.empresa)
    data = []
    for c in artigos:
        data.append({
            "codigo": c.id_artigo,
            "nome": c.nome,
            "descricao": c.descricao,
            "preco": c.preco,
            "taxa": c.taxa,
        })
    return JsonResponse(data, safe=False)


PRAZO_PAGAMENTO = {
    'Pronto Pagamento': 0,
    '30 DIAS': 30,
    '45 DIAS': 45,
    '60 DIAS': 60,
}


@login_required
@empresa_obrigatoria
def get_clientes(request):
    # Usamos prefetch_related em vez de select_related para evitar o INNER JOIN
    # que esconde clientes sem pagamento
    clientes = Cliente1.objects.filter(empresa=request.empresa).prefetch_related('pagamento')

    data = []
    for c in clientes:
        prazo_dias = 0

        try:
            if c.pagamento:
                nome_pagamento = c.pagamento.nome.strip()
                prazo_dias = PRAZO_PAGAMENTO.get(nome_pagamento, 0)
        except:
            prazo_dias = 0

        data.append({
            'id': c.id_cliente,
            'nome': c.nome,
            'prazo': prazo_dias,
        })

    return JsonResponse(data, safe=False)

@login_required
@empresa_obrigatoria
def adicionar_item(request, template_name='faturas/nova_fatura.html'):
    artigos = Artigo.objects.filter(empresa=request.empresa).order_by('descricao')

    # Debug: Imprime no terminal para veres se os dados existem
    print(f"--- Artigos da Empresa: {request.empresa} ---")
    for art in artigos:
        print(f"ID: {art.id_artigo} | Nome: {art.nome} | Preço: {art.preco} | Taxa: {art.taxa}")
    print(f"Total encontrado: {artigos.count()}")

    context = {
        'artigos': artigos,
    }
    return render(request, template_name, context)

from django.views.decorators.csrf import csrf_exempt
import json

@login_required
@empresa_obrigatoria
@csrf_exempt
def validar_linha(request):
    if request.method != "POST":
        return JsonResponse({"ok": False, "erros": ["Método inválido"]})

    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({"ok": False, "erros": ["Dados inválidos"]})

    erros = []

    tipo_doc = data.get("tipoDocumento", "FT")
    code = data.get("code")
    item = data.get("item")

    try:
        quantity = float(data.get("quantity", 0))
        price = float(data.get("price", 0))
        discount = float(data.get("discount", 0))
        tax = int(data.get("tax", 0))
    except (ValueError, TypeError):
        return JsonResponse({"ok": False, "erros": ["Formatos numéricos inválidos"]})

    motivo_iva0 = data.get("motivo_iva0")
    validacao_final = data.get("validacao_final", False)


    if not code:
        erros.append("Deves selecionar um artigo (código).")
    if not item:
        erros.append("A descrição do artigo não pode estar vazia.")

    if tipo_doc == "GT":
        if price < 0:
            erros.append("O preço na guia não pode ser negativo.")
    else:
        if price <= 0:
            erros.append("O preço deve ser um número maior que 0.")

    if quantity <= 0:
        erros.append("A quantidade deve ser um número maior que 0.")

    if discount < 0:
        erros.append("O desconto deve ser um número >= 0.")

    if tipo_doc == "GT":
        tax = 0
    else:
        # Apenas valida taxas para FT, FR, FS, NC
        if tax not in [0, 6, 13, 23]:
            erros.append("O IVA deve ser 0%, 6%, 13% ou 23%.")

    if tax == 0 and validacao_final:
        if not motivo_iva0:
            erros.append("Quando o IVA é 0%, é obrigatório indicar o motivo de isenção.")

    total_calculado = price * quantity
    if discount > total_calculado and price > 0:
        erros.append("O desconto não pode ser maior que o total da linha.")

    if erros:
        return JsonResponse({"ok": False, "erros": erros})

    return JsonResponse({"ok": True})

@login_required
@empresa_obrigatoria
def matriculas_dropdown(request):
    transportes = Transporte.objects.filter(empresa=request.empresa).order_by('descricao')  # ou outro filtro
    data = [{"descricao": t.descricao} for t in transportes]
    return JsonResponse(data, safe=False)

def obter_proximo_numero_final(tipo, serie, ano, empresa):
    with transaction.atomic():
        contador, created = DocumentoFinalizadoContador.objects.select_for_update().get_or_create(
            empresa=empresa,
            tipo=tipo,
            serie=serie,
            ano=ano,
            defaults={"ultimo_numero": 0}
        )

        numero = contador.ultimo_numero + 1
        contador.ultimo_numero = numero
        contador.save()

        return numero

from django.db import transaction
from django.views.decorators.csrf import csrf_exempt
from datetime import datetime, timedelta


@login_required
@empresa_obrigatoria
@csrf_exempt
def criar_documento_temp(request):
    if request.method != "POST":
        return JsonResponse({"erro": "Método inválido."}, status=405)

    cliente_id = request.POST.get("cliente")
    documento = request.POST.get("documento")
    data_emissao = request.POST.get("data_emissao")
    data_vencimento = request.POST.get("data_vencimento")

    if not all([cliente_id, documento, data_emissao, data_vencimento]):
        return JsonResponse({"erro": "Campos obrigatórios em falta."}, status=400)

    try:
        # Certifique-se de que os dados do cliente sejam válidos
        cliente = Cliente1.objects.select_related(
            "impostos"
        ).get(id_cliente=cliente_id,
              empresa=request.empresa
              )
    except Cliente1.DoesNotExist:
        return JsonResponse({"erro": "Cliente inválido."}, status=400)

    try:
        # Converte as datas para objetos datetime
        data_emissao = datetime.strptime(data_emissao, "%Y-%m-%d")  # Supondo que a data seja no formato "YYYY-MM-DD"
        data_vencimento = datetime.strptime(data_vencimento, "%Y-%m-%d")
    except ValueError:
        return JsonResponse({"erro": "Formato de data inválido. Use YYYY-MM-DD."}, status=400)

    ano = data_emissao.year

    with transaction.atomic():
        contador, created = DocumentoContador.objects.select_for_update().get_or_create(
            empresa=request.empresa,
            tipo=documento,
            ano=ano,
            defaults={"serie": "A", "ultimo_numero": 0}
        )

        # Incrementa o número do documento
        numero = contador.ultimo_numero + 1
        contador.ultimo_numero = numero
        contador.save()

        # Cria o documento temporário
        doc = DocumentoTemp.objects.create(
            tipo=documento,
            serie=contador.serie,
            numero=numero,
            ano=ano,
            cliente=cliente,
            data_emissao=data_emissao,
            data_vencimento=data_vencimento,
            valor_total=0,
            transporte=cliente.transporte,
            impostos=cliente.impostos,
            pagamento=cliente.pagamento,
            moeda_id=15,
            empresa=request.empresa
        )

    return JsonResponse({
        "id": doc.id,
        "tipo": doc.tipo,
        "serie": doc.serie,
        "numero": doc.numero,
        "ano": doc.ano,
    })

from .models import Empresa

@login_required
@empresa_obrigatoria
def obter_documento_temp(request, temp_id):
    try:
        doc = DocumentoTemp.objects.select_related(
            "cliente__transporte",
            "cliente__pagamento",
            "cliente__impostos",
            "moeda"
        ).get(id=temp_id,
              empresa=request.empresa
        )
    except DocumentoTemp.DoesNotExist:
        return JsonResponse({"erro": "Documento não encontrado"}, status=404)

    moedas = list(Moeda.objects.values("id", "codigo", "nome", "simbolo"))

    modalidades = list(Modalidade.objects.values("id_modalidade", "nome"))

    tax_reasons = list(TaxReason.objects.values("code", "description"))

    cliente_data = None
    if doc.cliente:
        cliente_data = {
            "id": doc.cliente.id_cliente,
            "codigo": doc.cliente.codigo,
            "nome": doc.cliente.nome,
            "morada1": doc.cliente.morada1,
            "morada2": doc.cliente.morada2,
            "codigo_postal": doc.cliente.codigo_postal,
            "pais": doc.cliente.pais,
            "concelho": doc.cliente.concelho,
            "contribuinte": doc.cliente.contribuinte,
            "email": doc.cliente.email,
            "transporte": doc.cliente.transporte.descricao if (
                        doc.cliente.transporte and doc.cliente.transporte.descricao) else None,
            "modalidade": {
                "id": doc.cliente.modalidade.id_modalidade,
                "nome": doc.cliente.modalidade.nome
            } if doc.cliente.modalidade else None,
            "impostos": doc.cliente.impostos.nome
        }
        if doc.cliente.impostos.nome.upper() == "IVA":
            cliente_data["local_descarga"] = doc.cliente.morada1
        else:
            cliente_data["local_descarga"] = ""

    artigos_qs = TempArtigos.objects.filter(id_temp=doc, empresa=request.empresa).select_related(
        "id_art",
        "motivo"
    )

    artigos = []

    for a in artigos_qs:
        artigos.append({
            "id": a.id_art.id_artigo,
            "tipo": a.tipo,
            "descricao": a.descricao,
            "quantidade": a.quantidade,
            "preco": float(a.preco),
            "desconto": float(a.desconto),
            "taxa": float(a.taxa),
            "total": float(a.total),
            "motivo": a.motivo.code if a.motivo else None,
            "motivo_descricao": a.motivo.description if a.motivo else None
        })

    empresa = request.empresa
    return JsonResponse({
        # --- DOCUMENTO ---
        "id": doc.id,
        "tipo": doc.tipo,
        "serie": doc.serie,
        "numero": doc.numero,
        "ano": doc.ano,
        "ordem_compra": doc.ordem_compra,
        "numero_compromisso": doc.numero_compromisso,
        "descricao": doc.descricao,
        "rodape": doc.rodape,
        "data_emissao": doc.data_emissao,
        "data_vencimento": doc.data_vencimento,
        "valor_total": doc.valor_total,
        "data_carga": doc.data_carga,
        "data_descarga": doc.data_descarga,
        "expedicao": doc.expedicao,
        "matricula": doc.transporte if doc.transporte else None,

        # --- CLIENTE ---
        "cliente": cliente_data,
        "modalidades": modalidades,

        # --- MOEDA ---
        "moeda": {
            "id": doc.moeda.id,
            "codigo": doc.moeda.codigo,
            "nome": doc.moeda.nome,
            "simbolo": doc.moeda.simbolo,
        },

        # lista de moedas disponíveis
        "moedas": moedas,

        # --- EMPRESA ---
        "empresa": {
            "nome": empresa.nome,
            "morada": empresa.morada,
            "codigo_postal": empresa.codigo_postal,
            "cidade": empresa.cidade,
            "pais": empresa.pais,
            "email": empresa.email,
            "nif": empresa.nif,
            "telefone": empresa.telefone,
            "local": empresa.local
        },

        "tax_reasons": tax_reasons,
        "artigos": artigos
    })

class Command(BaseCommand):
    help = "Apaga documentos temporários com tipo específico se tiverem mais de 7 dias sem alterações"

    def handle(self, *args, **kwargs):
        limite = timezone.now() - timedelta(days=7)
        documentos = DocumentoTemp.objects.filter(
            estado='Rascunho',
            atualizado_em__lt=limite,
        )
        count = documentos.count()
        documentos.delete()
        self.stdout.write(self.style.SUCCESS(f'Apagados {count} documentos antigos.'))

@login_required
@empresa_obrigatoria
def apagar_documento(request):
    if request.method == "POST":
        import json
        data = json.loads(request.body)
        doc_id = data.get("id")
        doc_tipo = data.get("tipo")

        if not doc_id:
            return JsonResponse({"success": False, "error": "ID do documento não fornecido."}, status=400)

        try:
            if doc_tipo == "TEMP":
                documento = DocumentoTemp.objects.filter(id=doc_id, empresa=request.empresa).first()
                if not documento:
                    return JsonResponse({"success": False, "error": "Documento temporário não encontrado."}, status=404)
                documento.delete()

            elif doc_tipo == "NC":
                documento = DocumentoFinalizado.objects.filter(id=doc_id, tipo='NC', empresa=request.empresa).first()
                if not documento:
                    return JsonResponse({"success": False, "error": "Nota de Crédito não encontrada."}, status=404)
                documento.estado = 'Anulado'
                documento.save()

            else:
                return JsonResponse({"success": False, "error": "Tipo de documento não suportado."}, status=400)

            return JsonResponse({"success": True, "message": "Documento processado com sucesso."})

        except Exception as e:
            return JsonResponse({"success": False, "error": f"Erro ao processar documento: {str(e)}"}, status=400)

    return JsonResponse({"success": False, "error": "Método inválido."}, status=405)

import re
@login_required
@empresa_obrigatoria
def validar_ordem_compra(ordem_compra):
    """ Valida se a ordem de compra segue o formato esperado (ex: 'PO12345'). """
    if ordem_compra:
        if not re.match(r'^[A-Za-z]{2,4}\d{1,6}$', ordem_compra):  # Exemplo de padrão alfanumérico (PO12345)
            raise ValidationError("Ordem de compra inválida. O formato esperado é 'PO12345'.")
    return ordem_compra

@login_required
@empresa_obrigatoria
def validar_numero_compromisso(numero_compromisso):
    """ Valida se o número de compromisso é um número válido e dentro do limite esperado. """
    if numero_compromisso:
        if not numero_compromisso.isdigit():  # Verifica se o número de compromisso é numérico
            raise ValidationError("Número de compromisso inválido. Apenas números são permitidos.")
        if len(numero_compromisso) < 6 or len(numero_compromisso) > 12:
            raise ValidationError("Número de compromisso inválido. O comprimento deve ser entre 6 e 12 caracteres.")
    return numero_compromisso

@login_required
@empresa_obrigatoria
def validar_texto_longo(campo_texto, limite=500):
    """ Verifica se o texto ultrapassa o limite de caracteres e o corta se necessário. """
    if campo_texto:
        if len(campo_texto) > limite:
            raise ValidationError(f"O texto não pode exceder {limite} caracteres.")
        return campo_texto
    return campo_texto

import logging

logger = logging.getLogger(__name__)
from django.core.exceptions import ValidationError
from decimal import Decimal, ROUND_HALF_UP

@login_required
@empresa_obrigatoria
def atualizar_documento(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            temp_id = data.get("temp_id")
            if not temp_id:
                return JsonResponse({"success": False, "error": "temp_id não fornecido"}, status=400)

            documento = DocumentoTemp.objects.filter(id=temp_id, empresa=request.empresa).first()
            if not documento:
                return JsonResponse({"success": False, "error": "Documento não encontrado"}, status=404)

            # Atualizar campos do documento
            documento.data_emissao = data.get("data_emissao") or documento.data_emissao
            documento.data_vencimento = data.get("data_vencimento") or documento.data_vencimento
            documento.ordem_compra = data.get("ordem_compra") or documento.ordem_compra
            documento.numero_compromisso = data.get("numero_comp") or documento.numero_compromisso
            documento.descricao = data.get("descricao") or documento.descricao
            documento.rodape = data.get("rodape") or documento.rodape
            documento.local_carga = data.get("local_carga") or documento.local_carga
            documento.local_descarga = data.get("local_descarga") or documento.local_descarga

            ordem_compra = data.get("ordem_compra")
            if ordem_compra:
                documento.ordem_compra = validar_ordem_compra(ordem_compra)

            numero_compromisso = data.get("numero_comp")
            if numero_compromisso:
                documento.numero_compromisso = validar_numero_compromisso(numero_compromisso)

            descricao = data.get("descricao")
            if descricao:
                documento.descricao = validar_texto_longo(descricao, limite=500)

            rodape = data.get("rodape")
            if rodape:
                documento.rodape = validar_texto_longo(rodape, limite=200)
            try:
                data_emissao = data.get("data_emissao")
                if data_emissao:
                    documento.data_emissao = datetime.strptime(data_emissao, "%d/%m/%Y")

                    # Verificar se a data de emissão é maior ou igual à data atual
                    if documento.data_emissao.date() < timezone.now().date():
                        return JsonResponse(
                            {"success": False, "error": "A data de emissão não pode ser anterior à data atual."},
                            status=400)

                local_carga = data.get("local_carga")
                if local_carga:
                    if len(local_carga) > 255:  # Limite de 255 caracteres para local de carga
                        raise ValidationError("O local de carga não pode ter mais que 255 caracteres.")

                local_descarga = data.get("local_descarga")
                if local_descarga:
                    if len(local_descarga) > 255:  # Limite de 255 caracteres para local de descarga
                        raise ValidationError("O local de descarga não pode ter mais que 255 caracteres.")

                data_vencimento = data.get("data_vencimento")
                if data_vencimento:
                    documento.data_vencimento = datetime.strptime(data_vencimento, "%d/%m/%Y")

                    # Verificar se a data de vencimento é maior ou igual à data atual
                    if documento.data_vencimento.date() < timezone.now().date():
                        return JsonResponse(
                            {"success": False, "error": "A data de vencimento não pode ser anterior à data atual."},
                            status=400)

                    # Validar se a data de vencimento não é anterior à data de emissão
                    if documento.data_vencimento.date() < documento.data_emissao.date():
                        return JsonResponse({"success": False,
                                             "error": "A data de vencimento não pode ser anterior à data de emissão."},
                                            status=400)

                data_carga = data.get("data_carga")
                if data_carga:
                    # converter string para datetime
                    dt_carga_naive = datetime.strptime(data_carga, "%d/%m/%Y %H:%M")

                    # tornar timezone-aware
                    documento.data_carga = timezone.make_aware(dt_carga_naive, timezone.get_current_timezone())

                    if documento.data_carga.date() < timezone.now().date():
                        return JsonResponse(
                            {"success": False, "error": "A data de carga não pode ser anterior à data atual."},
                            status=400
                        )

                # Data de descarga
                data_descarga = data.get("data_descarga")
                if data_descarga:
                    dt_descarga_naive = datetime.strptime(data_descarga, "%d/%m/%Y %H:%M")
                    documento.data_descarga = timezone.make_aware(dt_descarga_naive, timezone.get_current_timezone())

                    if documento.data_descarga < timezone.now() :
                        return JsonResponse(
                            {"success": False, "error": "A data de descarga não pode ser anterior à data atual."},
                            status=400
                        )

                    if documento.data_descarga < documento.data_carga:
                        return JsonResponse(
                            {"success": False, "error": "A data de descarga não pode ser anterior à data de carga."},
                            status=400
                        )
            except ValueError as e:
                return JsonResponse({"success": False, "error": f"Erro no formato da data: {str(e)}"}, status=400)

            if documento.data_carga and documento.data_emissao:
                if documento.data_carga.date() < documento.data_emissao.date():
                    return JsonResponse(
                        {"success": False, "error": "A data de carga não pode ser anterior à data de emissão."},
                        status=400
                    )
            documento.expedicao = data.get("expedicao") or documento.expedicao

            expedicao = data.get("expedicao")
            if expedicao:
                if len(expedicao) > 255:  # Limite de 255 caracteres para local de descarga
                    raise ValidationError("O modo de expedição não pode ter mais que 255 caracteres.")

            cliente_id = data.get("cliente_id")
            if not cliente_id:
                return JsonResponse(
                    {"success": False, "error": "Cliente obrigatório."},
                    status=400
                )
            if cliente_id:
                cliente_obj = Cliente1.objects.filter(id_cliente=cliente_id, empresa=request.empresa).first()
                if not cliente_obj:
                    return JsonResponse({"success": False, "error": f"Cliente {cliente_id} não encontrado."}, status=400)
                documento.cliente_id = cliente_obj.id_cliente

            metodo_pagamento = data.get("metodo_pagamento")
            if metodo_pagamento:
                pagamento_obj = Modalidade.objects.filter(nome=metodo_pagamento).first()
                if not pagamento_obj:
                    return JsonResponse({"success": False, "error": f"Modalidade de pagamento {metodo_pagamento} não encontrada."}, status=400)
                documento.pagamento_id = pagamento_obj.id_modalidade

            moeda_id = data.get("moeda")
            if moeda_id:
                moeda_obj = Moeda.objects.filter(id=moeda_id).first()
                if not moeda_obj:
                    return JsonResponse({"success": False, "error": f"Moeda {moeda_id} não encontrada."}, status=400)
                documento.moeda_id = moeda_obj.id

            transporte_descricao = data.get("matricula")

            if transporte_descricao:
                # 1. Validação de tamanho (limite do CharField na BD)
                if len(transporte_descricao) > 255:
                    return JsonResponse({
                        "success": False,
                        "error": "A matrícula não pode ter mais de 255 caracteres."
                    }, status=400)


            TempArtigos.objects.filter(id_temp=documento.id, empresa=request.empresa).delete()

            # Criar novos artigos
            artigos_data = data.get("artigos", [])
            if not artigos_data:
                return JsonResponse(
                    {"success": False, "error": "O documento deve conter pelo menos um artigo."},
                    status=400
                )
            valor_total = Decimal("0.00")
            if artigos_data:
                for artigo_data in artigos_data:
                    try:
                        # Aqui pode ocorrer um erro se 'codigo' não for fornecido ou o formato de dados estiver errado
                        artigo_codigo = artigo_data.get("codigo")
                        if not artigo_codigo:
                            logger.error(f"Erro: Código do artigo não fornecido para os dados {artigo_data}")
                            return JsonResponse({"success": False, "error": "Código do artigo não fornecido"},
                                                status=400)

                        artigo_obj = Artigo.objects.filter(id_artigo=artigo_codigo, empresa=request.empresa).first()
                        if not artigo_obj:
                            logger.error(f"Erro: Artigo com código {artigo_codigo} não encontrado.")
                            return JsonResponse(
                                {"success": False, "error": f"Artigo com código {artigo_codigo} não encontrado."},
                                status=400)

                        # Continuar a criação do artigo, mais validações podem ser feitas aqui
                        tipo_artigo = artigo_obj.tipo
                        preco = Decimal(artigo_data.get("preco", "0").replace(",", "."))
                        if preco <= 0:
                            return JsonResponse(
                                {"success": False, "error": "O preço deve ser maior que zero."},
                                status=400
                            )
                        desconto = Decimal(artigo_data.get("desconto", "0").replace(",", "."))
                        if desconto < 0:
                            return JsonResponse(
                                {"success": False, "error": "O desconto não pode ser negativo."},
                                status=400
                            )

                        iva = Decimal(artigo_data.get("iva", "0").replace(",", "."))
                        TAXAS_VALIDAS = {0, 6, 13, 23}

                        if iva not in TAXAS_VALIDAS:
                            return JsonResponse(
                                {"success": False, "error": "Taxa de IVA inválida."},
                                status=400
                            )
                        total = Decimal(artigo_data.get("total", "0").replace(",", "."))


                        quantidade_raw = artigo_data.get("quantidade", 0)

                        try:
                            quantidade = int(quantidade_raw)
                        except (ValueError, TypeError):
                            return JsonResponse(
                                {"success": False, "error": "Quantidade inválida."},
                                status=400
                            )

                        if quantidade < 1 or quantidade > 10000:
                            return JsonResponse(
                                {"success": False, "error": "A quantidade deve estar entre 1 e 10000."},
                                status=400
                            )
                        if desconto > preco * quantidade:
                            return JsonResponse(
                                {"success": False, "error": "O desconto não pode ser superior ao valor da linha."},
                                status=400
                            )
                        descricao = artigo_data.get("descricao", "").strip()

                        if not descricao:
                            return JsonResponse(
                                {"success": False, "error": "A descrição do artigo é obrigatória."},
                                status=400
                            )

                        if len(descricao) > 255:
                            return JsonResponse(
                                {"success": False, "error": "A descrição não pode exceder 255 caracteres."},
                                status=400
                            )

                        subtotal = (preco * quantidade - desconto).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
                        total_com_iva = (subtotal * (1 + iva / 100)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)

                        total_recebido = total.quantize(
                            Decimal("0.01"), rounding=ROUND_HALF_UP
                        )

                        valor_total += total_com_iva

                        if subtotal != total_recebido:
                            return JsonResponse(
                                {
                                    "success": False,
                                    "error": f"Total inválido. Esperado: {subtotal}, Recebido: {total_recebido}"
                                },
                                status=400
                            )

                        motivo_tax = artigo_data.get("motivo")

                        if iva == 0:
                            if not motivo_tax:
                                return JsonResponse(
                                    {"success": False, "error": "Motivo de IVA obrigatório quando a taxa é 0%."},
                                    status=400
                                )

                            tax_reason = TaxReason.objects.filter(code=motivo_tax).first()
                            if not tax_reason:
                                logger.error(f"Erro: Motivo IVA '{motivo_tax}' não encontrado.")
                                return JsonResponse(
                                    {"success": False, "error": f"Motivo IVA '{motivo_tax}' inválido."},
                                    status=400
                                )

                        else:
                            if motivo_tax:
                                return JsonResponse(
                                    {"success": False, "error": "Motivo de IVA só é permitido quando a taxa é 0%."},
                                    status=400
                                )

                            tax_reason = None

                        if documento.tipo == 'FS':
                            limite_consumidor_final = Decimal("1000.00")
                            limite_sujeito_passivo = Decimal("100.00")

                            tem_nif = False
                            if documento.cliente and documento.cliente.contribuinte:
                                if str(documento.cliente.contribuinte) != "999999990":
                                    tem_nif = True

                            if tem_nif and valor_total > limite_sujeito_passivo:
                                return JsonResponse({
                                    "success": False,
                                    "error": f"Fatura Simplificada para empresas não pode exceder {limite_sujeito_passivo}€. Use Fatura (FT)."
                                }, status=400)

                            if valor_total > limite_consumidor_final:
                                return JsonResponse({
                                    "success": False,
                                    "error": f"Fatura Simplificada não pode exceder {limite_consumidor_final}€. Use Fatura (FT)."
                                }, status=400)

                        TempArtigos.objects.create(
                            id_temp=documento,
                            id_art=artigo_obj,
                            tipo=tipo_artigo,
                            descricao=descricao,
                            quantidade=quantidade,
                            preco=preco,
                            desconto=desconto,
                            taxa=iva,
                            total=subtotal,
                            motivo_id=tax_reason.id if tax_reason else None,
                            empresa=request.empresa
                        )

                    except Exception as e:
                        # Log de erro específico do artigo
                        logger.error(f"Erro ao processar artigo {artigo_data}: {str(e)}")
                        return JsonResponse({"success": False, "error": f"Erro ao processar artigo: {str(e)}"},
                                            status=400)
            if documento.tipo == 'NC':
                valor_total = -abs(valor_total)

            documento.valor_total = valor_total
            documento.empresa = request.empresa
            documento.save()

            return JsonResponse({
                "success": True,
                "message": "Documento atualizado com sucesso!",
                "documento_id": documento.id
            })

        except Exception as e:
            logger.error(f"Erro inesperado ao atualizar documento: {str(e)}", exc_info=True)  # Log detalhado do erro
            return JsonResponse({"success": False, "error": f"Erro inesperado: {str(e)}"}, status=400)


from .models import DocumentoTemp, TempArtigos

@login_required
@empresa_obrigatoria
def editar_fatura(request):
    """
    View para editar uma fatura temporária.
    Recebe o temp_id (id do DocumentoTemp) na URL.
    Retorna todos os dados do documento e todos os artigos relacionados.
    """

    # Pega o ID do documento da URL
    temp_id = request.GET.get('temp_id')

    # Busca o documento temporário
    documento = get_object_or_404(DocumentoTemp, id=temp_id, empresa=request.empresa)

    artigos = TempArtigos.objects.filter(id_temp=documento, empresa=request.empresa).select_related('id_art', 'motivo')

    artigos_novos = Artigo.objects.filter(empresa=request.empresa)

    # Renderiza o template, passando documento e artigos
    return render(request, 'faturas/editar_faturas.html', {
        'documento': documento,
        'artigos': artigos,
        'artigos_novos': artigos_novos
    })

def gerar_codigo_at(documento):
    """
    Gera um código único da AT para o documento.
    Formato: TIPO-SERIE-ANO-NUMERO
    Exemplo: FT-A-2026-000123
    """
    return f"{documento.tipo}-{documento.serie}-{documento.ano}-{str(documento.numero).zfill(6)}"

from django.conf import settings

from docx.oxml.ns import qn

from docxtpl import DocxTemplate
from collections import defaultdict
from copy import deepcopy
from docx.oxml import OxmlElement
import copy
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph

def remover_bordas(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = OxmlElement('w:tcBorders')

    for lado in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        elem = OxmlElement(f'w:{lado}')
        elem.set(qn('w:val'), 'nil')  # sem borda
        tcBorders.append(elem)

    tcPr.append(tcBorders)


import qrcode
from io import BytesIO
from docx.shared import Mm
from docxtpl import InlineImage

def gerar_qr(documento: DocumentoFinalizado, docx_obj):
    """
    Gera o QR Code oficial da fatura e retorna como InlineImage para inserir no docx.
    """
    # Montar dados do QR Code
    qr_data = (
        f"A:{documento.empresa_contribuinte}*"
        f"B:{documento.cliente_contribuinte or '999999990'}*"
        f"C:{documento.cliente_pais or 'PT'}*"
        f"D:{documento.tipo}*"
        f"E:N*"
        f"F:{documento.data_emissao.strftime('%Y%m%d')}*"
        f"G:{documento.tipo} {documento.serie}/{documento.numero}*"
        f"H:{documento.valor_total:.2f}*"
        f"P:{documento.codigo_at_tributaria or ''}"
    )

    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=4,
        border=1,
    )
    qr.add_data(qr_data)
    qr.make(fit=True)

    img = qr.make_image(fill_color="black", back_color="white")

    # Salvar em buffer
    qr_buffer = BytesIO()
    img.save(qr_buffer, format="PNG")
    qr_buffer.seek(0)

    # Criar InlineImage do docxtpl
    qr_inline = InlineImage(docx_obj, qr_buffer, width=Mm(30))  # 30 mm de largura
    return qr_inline

from docx.shared import Pt, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def encontrar_tabela_por_conteudo(doc, fragmento_texto):
    fragmento = fragmento_texto.lower().strip()
    for i, tabela in enumerate(doc.tables):
        for row in tabela.rows:
            for cell in row.cells:
                if fragmento in cell.text.lower():
                    return tabela
    return None

@login_required
@empresa_obrigatoria
def gerar_word_fatura(request, documento_final, via="original"):

    modelo_path = os.path.join(
        settings.BASE_DIR,
        "Faturamento/docs/Documento.docx"
    )
    artigos = FinArtigos.objects.filter(id_final=documento_final, empresa=request.empresa)

    mapa_iva = defaultdict(lambda: {"base": Decimal("0.00"), "valor": Decimal("0.00")})
    subtotal = Decimal("0.00")
    for art in artigos:
        taxa = Decimal(str(art.taxa))
        base = art.quantidade * art.preco
        base -= art.desconto or Decimal("0.00")
        subtotal += base
        valor_iva = base * taxa / Decimal("100")
        mapa_iva[taxa]["base"] += base
        mapa_iva[taxa]["valor"] += valor_iva

    context = {
        "documento_nome": str(documento_final),
        "tipo_documento": documento_final.tipo,
        "tipo_fatura": documento_final.get_tipo_display(),
        "nome_empresa": documento_final.empresa_nome,
        "morada_empresa": documento_final.empresa_morada,
        "postal_empresa": documento_final.empresa_codigo_postal,
        "concelho_empresa": documento_final.empresa_cidade,
        "nif_empresa": documento_final.empresa_contribuinte,

        "nome_cliente": documento_final.cliente_nome,
        "morada_cliente": documento_final.cliente_morada1,
        "postal_cliente": documento_final.cliente_codigo_postal,
        "concelho_cliente": documento_final.cliente_concelho,
        "nif_cliente": documento_final.cliente_contribuinte,

        "data_emissao": documento_final.data_emissao.strftime("%d/%m/%Y"),
        "data_vencimento": documento_final.data_vencimento.strftime("%d/%m/%Y"),


        "cod_at": documento_final.codigo_at_tributaria,
        "tipo_moeda": documento_final.moeda_simbolo,
        "ordem_compra": documento_final.ordem_compra,
        "numero_compromisso": documento_final.numero_compromisso,
        "subtotal": subtotal,
        "total_fatura": documento_final.valor_total,
        "metodo_pagamento": (
            documento_final.modalidade_nome
            if documento_final.modalidade_nome else None
        ),
        "local_carga": documento_final.local_carga if documento_final.local_carga else "",
        "local_descarga": documento_final.local_descarga if documento_final.local_descarga else "",

        # Formatação de datas (caso existam)
        "data_carga": (
            documento_final.data_carga.strftime("%d/%m/%Y")
            if documento_final.data_carga else ""
        ),
        "data_descarga": (
            documento_final.data_descarga.strftime("%d/%m/%Y")
            if documento_final.data_descarga else ""
        ),

        "expedição": documento_final.expedicao if documento_final.expedicao else "",
        "matricula": documento_final.transporte_descricao if documento_final.transporte_descricao else "",    }


    if documento_final.descricao:
        context["descricao"] = documento_final.descricao

    if documento_final.rodape:
        context["rodape"] = documento_final.rodape

    # =============================
    # ARTIGOS
    # =============================

    artigos = FinArtigos.objects.filter(id_final=documento_final, empresa=request.empresa)

    vias_nomes = {
        "original": ["Original"],
        "duplicado": ["Original", "Duplicado"],
        "triplicado": ["Original", "Duplicado", "Triplicado"]
    }.get(via, ["Original"])
    documentos_vias = []


    for nome_via in vias_nomes:
        doc_via = DocxTemplate(modelo_path)


        # Criar contexto específico para esta via
        ctx = deepcopy(context)
        ctx["ficheiro"] = nome_via
        ctx["num_folha"] = 1
        ctx["total_folhas"] = 1
        ctx["qr_code"] = gerar_qr(documento_final, doc_via)

        # Renderiza as tags (incluindo as do cabeçalho e o IF da carga/descarga)
        doc_via.render(ctx)

        tabela_artigos = encontrar_tabela_por_conteudo(doc_via, "lista")

        for idx, a in enumerate(artigos, start=1):
            row = tabela_artigos.add_row().cells
            cor_hex = "F2F2F2" if idx % 2 != 0 else "E6E6E6"

            textos = [
                a.descricao,
                str(a.quantidade),
                f"{a.preco:.2f}",
                f"{a.desconto:.2f}",
                f"{a.taxa:.0f} *{idx}" if a.motivo else f"{a.taxa:.0f}",
                f"{a.total:.2f}"
            ]

            for i, cell in enumerate(row):
                # CORREÇÃO: Gerar um novo elemento de shading para CADA célula
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{cor_hex}"/>')
                cell._tc.get_or_add_tcPr().append(shading_elm)

                # Se a tua função remover_bordas mexer no XML da célula,
                # garante que ela não apaga o tcPr que acabaste de criar
                remover_bordas(cell)

                par = cell.paragraphs[0]
                # Limpar runs existentes
                for run_existente in par.runs:
                    par._element.remove(run_existente._element)

                par.paragraph_format.space_before = Pt(6)
                par.paragraph_format.space_after = Pt(6)
                par.paragraph_format.line_spacing = 1.2

                run = par.add_run(textos[i])
                run.font.name = "Geologica"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Geologica')
                run.font.size = Pt(9)
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # =============================
        # RESUMO DE IVA
        # ============================
        tabela_resumo = encontrar_tabela_por_conteudo(doc_via, "subtotal")

        # 1️⃣ Localizar o elemento XML (TR) da linha do Subtotal
        linha_referencia_tr = None
        for linha in tabela_resumo.rows:
            if "subtotal" in linha.cells[0].text.lower():
                linha_referencia_tr = linha._tr
                break

        # Se não encontrar o subtotal, usamos a primeira linha como base
        if linha_referencia_tr is None:
            linha_referencia_tr = tabela_resumo.rows[0]._tr

        # 2️⃣ Inserir as linhas de IVA
        for taxa, dados in sorted(mapa_iva.items(), reverse=True):
            # Criamos um novo elemento de linha (TR)
            # Em vez de add_row(), criamos um elemento vazio para ter controle total
            nova_linha_tr = copy.deepcopy(linha_referencia_tr)

            # Limpamos o texto das células da linha clonada
            for cell_xml in nova_linha_tr.xpath('.//w:t'):
                cell_xml.text = ""

            # Inserimos a nova linha logo após a linha de referência no XML
            linha_referencia_tr.addnext(nova_linha_tr)

            # Agora acessamos essa linha como um objeto do python-docx para formatar
            from docx.table import _Row
            nova_linha = _Row(nova_linha_tr, tabela_resumo)

            # Preencher os dados (assumindo que a tabela tem pelo menos 2 colunas)
            textos = [
                f"IVA {taxa:.0f}% (Incidência {dados['base']:.2f})",
                f"{dados['valor']:.2f} {documento_final.moeda_simbolo}"
            ]

            for j, texto in enumerate(textos):
                celula = nova_linha.cells[j]
                # Limpar parágrafos existentes e adicionar o novo
                p = celula.paragraphs[0]
                p.clear()
                run = p.add_run(texto)
                run.font.name = "Geologica"
                run.font.size = Pt(9)

            # IMPORTANTE: A próxima taxa deve vir depois desta que acabámos de criar
            # Por isso, a nova linha passa a ser a nossa referência
            linha_referencia_tr = nova_linha_tr

        motivos_agregados = defaultdict(list)

        for idx, art in enumerate(artigos, start=1):
            if art.motivo:
                motivos_agregados[art.motivo.description].append(idx)

        # --- 2️⃣ Inserir título se houver motivos ---
        tabela_xml = tabela_resumo._tbl
        parent = tabela_xml.getparent()

        from docx.oxml import OxmlElement
        if motivos_agregados:
            # 1️⃣ Criar parágrafo do título

            p_titulo_xml = OxmlElement("w:p")
            tabela_xml.addnext(p_titulo_xml)
            p_titulo = Paragraph(p_titulo_xml, tabela_resumo._parent)
            p_titulo.paragraph_format.space_before = Pt(24)

            run_titulo = p_titulo.add_run("Condições de Enquadramento de IVA:")
            run_titulo.bold = True
            run_titulo.font.name = "Geologica"
            run_titulo._element.rPr.rFonts.set(qn('w:eastAsia'), 'Geologica')
            run_titulo.font.size = Pt(10)
            p_titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT

            ultimo_elemento = p_titulo_xml

            # 2️⃣ Inserir linhas agrupadas
            for descricao, indices in motivos_agregados.items():
                p_xml = OxmlElement("w:p")
                ultimo_elemento.addnext(p_xml)
                p = Paragraph(p_xml, tabela_resumo._parent)

                numeros = "".join([f"(*{i})" for i in indices])
                run = p.add_run(f"- {numeros} {descricao}")
                run.font.name = "Geologica"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Geologica')
                run.font.size = Pt(9)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                ultimo_elemento = p_xml

        documentos_vias.append(doc_via)

    doc_final = documentos_vias[0]

    for extra_doc in documentos_vias[1:]:
        # Pegamos todos os elementos do corpo da via extra
        elementos_para_copiar = [
            deepcopy(el) for el in extra_doc.element.body
            if not isinstance(el, docx.oxml.section.CT_SectPr)
        ]

        if elementos_para_copiar:
            # 1. Forçamos uma quebra de página APENAS no primeiro parágrafo da nova via
            # Isso garante que a nova via começa no topo, sem criar páginas extras vazias
            primeiro_elemento = elementos_para_copiar[0]

            # Se o primeiro elemento for um parágrafo (<w:p>), inserimos a quebra lá
            if primeiro_elemento.tag.endswith('p'):
                run = OxmlElement('w:r')
                br = OxmlElement('w:br')
                br.set(qn('w:type'), 'page')
                run.append(br)
                primeiro_elemento.insert(0, run)
            else:
                # Se o primeiro elemento for uma tabela, criamos um parágrafo de quebra antes
                p = OxmlElement('w:p')
                r = OxmlElement('w:r')
                br = OxmlElement('w:br')
                br.set(qn('w:type'), 'page')
                r.append(br)
                p.append(r)
                doc_final.element.body.append(p)

            # 2. Adicionamos todos os elementos ao documento final
            for el in elementos_para_copiar:
                doc_final.element.body.append(el)

    buffer = BytesIO()
    doc_final.save(buffer)
    buffer.seek(0)

    return buffer


import sys
import subprocess
import os
import tempfile
import shutil
def converter_word_para_pdf(word_buffer):
    temp_dir = tempfile.mkdtemp()
    word_path = os.path.join(temp_dir, "input.docx")

    with open(word_path, "wb") as f:
        f.write(word_buffer.getbuffer())

    try:
        if sys.platform == "win32":
            # SE ESTIVER NO TEU PC (WINDOWS)
            # Precisas de ter o LibreOffice instalado.
            # O caminho padrão costuma ser este:
            libreoffice_path = r'C:\Program Files\LibreOffice\program\soffice.exe'
        else:
            # SE ESTIVER ONLINE (LINUX)
            libreoffice_path = 'libreoffice'

        subprocess.run([
            libreoffice_path,
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', temp_dir,
            word_path
        ], check=True)

        pdf_path = os.path.join(temp_dir, "input.pdf")
        with open(pdf_path, "rb") as f:
            return f.read()

    finally:
        shutil.rmtree(temp_dir)

@login_required
@empresa_obrigatoria
def gerar_pdf_fatura(request, documento_id):
    # Pega o documento
    documento = DocumentoFinalizado.objects.get(id=documento_id, empresa=request.empresa)

    via = request.GET.get("via", "original")

    word_buffer = gerar_word_fatura(request, documento, via)
    pdf_bytes = converter_word_para_pdf(word_buffer)

    # Retorna PDF como resposta HTTP
    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    response['Content-Disposition'] = f'inline; filename="Fatura_{documento.id}.pdf"'
    return response

from django.http import HttpResponse
@login_required
@empresa_obrigatoria
def finalizar_documento(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            temp_id = data.get("temp_id")
            if not temp_id:
                return JsonResponse({"success": False, "error": "temp_id não fornecido"}, status=400)

            documento = DocumentoTemp.objects.filter(id=temp_id, empresa=request.empresa).first()
            if not documento:
                return JsonResponse({"success": False, "error": "Documento não encontrado"}, status=404)

            documento.data_emissao = data.get("data_emissao") or documento.data_emissao
            documento.data_vencimento = data.get("data_vencimento") or documento.data_vencimento
            documento.ordem_compra = data.get("ordem_compra") or documento.ordem_compra
            documento.numero_compromisso = data.get("numero_comp") or documento.numero_compromisso
            documento.descricao = data.get("descricao") or documento.descricao
            documento.rodape = data.get("rodape") or documento.rodape
            documento.local_carga = data.get("local_carga") or documento.local_carga
            documento.local_descarga = data.get("local_descarga") or documento.local_descarga

            ordem_compra = data.get("ordem_compra")
            if ordem_compra:
                documento.ordem_compra = validar_ordem_compra(ordem_compra)

            numero_compromisso = data.get("numero_comp")
            if numero_compromisso:
                documento.numero_compromisso = validar_numero_compromisso(numero_compromisso)

            descricao = data.get("descricao")
            if descricao:
                documento.descricao = validar_texto_longo(descricao, limite=200)

            rodape = data.get("rodape")
            if rodape:
                documento.rodape = validar_texto_longo(rodape, limite=200)
            try:
                data_emissao = data.get("data_emissao")
                if data_emissao:
                    documento.data_emissao = datetime.strptime(data_emissao, "%d/%m/%Y")

                    # Verificar se a data de emissão é maior ou igual à data atual
                    if documento.data_emissao.date() < timezone.now().date():
                        return JsonResponse(
                            {"success": False, "error": "A data de emissão não pode ser anterior à data atual."},
                            status=400)

                tipo = documento.tipo

                local_carga = data.get("local_carga")
                local_descarga = data.get("local_descarga")

                if tipo in ("FR", "FT"):
                    if len(local_carga) > 255:
                        raise ValidationError("O local de carga não pode ter mais que 255 caracteres.")

                    if len(local_descarga) > 255:
                        raise ValidationError("O local de descarga não pode ter mais que 255 caracteres.")

                data_vencimento = data.get("data_vencimento")
                if data_vencimento:
                    documento.data_vencimento = datetime.strptime(data_vencimento, "%d/%m/%Y")

                    # Verificar se a data de vencimento é maior ou igual à data atual
                    if documento.data_vencimento.date() < timezone.now().date():
                        return JsonResponse(
                            {"success": False, "error": "A data de vencimento não pode ser anterior à data atual."},
                            status=400)

                    # Validar se a data de vencimento não é anterior à data de emissão
                    if documento.data_vencimento.date() < documento.data_emissao.date():
                        return JsonResponse({"success": False,
                                             "error": "A data de vencimento não pode ser anterior à data de emissão."},
                                            status=400)

                data_carga_str = data.get("data_carga")
                data_descarga_str = data.get("data_descarga")

                # ===== DATA DE CARGA =====
                if data_carga_str:
                    dt_carga_naive = datetime.strptime(data_carga_str, "%d/%m/%Y %H:%M")
                    data_carga = timezone.make_aware(
                        dt_carga_naive,
                        timezone.get_current_timezone()
                    )

                    if data_carga.date() < timezone.now().date():
                        return JsonResponse(
                            {"success": False, "error": "A data de carga não pode ser anterior à data atual."},
                            status=400
                        )

                    documento.data_carga = data_carga

                # ===== DATA DE DESCARGA =====
                if data_descarga_str:
                    dt_descarga_naive = datetime.strptime(data_descarga_str, "%d/%m/%Y %H:%M")
                    data_descarga = timezone.make_aware(
                        dt_descarga_naive,
                        timezone.get_current_timezone()
                    )

                    if data_descarga < timezone.now():
                        return JsonResponse(
                            {"success": False, "error": "A data de descarga não pode ser anterior à data atual."},
                            status=400
                        )

                    # Só comparar se existir data_carga
                    if data_carga_str and data_descarga < documento.data_carga:
                        return JsonResponse(
                            {"success": False, "error": "A data de descarga não pode ser anterior à data de carga."},
                            status=400
                        )

                    documento.data_descarga = data_descarga
            except ValueError as e:
                return JsonResponse({"success": False, "error": f"Erro no formato da data: {str(e)}"}, status=400)

            if documento.data_carga and documento.data_emissao:
                if documento.data_carga.date() < documento.data_emissao.date():
                    return JsonResponse(
                        {"success": False, "error": "A data de carga não pode ser anterior à data de emissão."},
                        status=400
                    )
            documento.expedicao = data.get("expedicao") or documento.expedicao

            expedicao = data.get("expedicao")
            if expedicao:
                if len(expedicao) > 255:  # Limite de 255 caracteres para local de descarga
                    raise ValidationError("O modo de expedição não pode ter mais que 255 caracteres.")

            cliente_id = data.get("cliente_id")
            if not cliente_id:
                return JsonResponse(
                    {"success": False, "error": "Cliente obrigatório."},
                    status=400
                )
            if cliente_id:
                cliente_obj = Cliente1.objects.filter(id_cliente=cliente_id, empresa=request.empresa).first()
                if not cliente_obj:
                    return JsonResponse({"success": False, "error": f"Cliente {cliente_id} não encontrado."}, status=400)
                documento.cliente_id = cliente_obj.id_cliente

            pagamento_obj = ""
            metodo_pagamento = data.get("metodo_pagamento")

            if metodo_pagamento:
                pagamento_obj = Modalidade.objects.filter(nome=metodo_pagamento).first()
                if pagamento_obj:
                    documento.pagamento_id = pagamento_obj.id_modalidade

            moeda_obj = None
            moeda_id = data.get("moeda")
            if moeda_id:
                moeda_obj = Moeda.objects.filter(id=moeda_id).first()
                if not moeda_obj:
                    return JsonResponse({"success": False, "error": f"Moeda {moeda_id} não encontrada."}, status=400)
                documento.moeda_id = moeda_obj.id

            transporte_descricao = data.get("matricula")
            if transporte_descricao:
                if len(transporte_descricao) > 255:
                    return JsonResponse({
                        "success": False,
                        "error": "A matrícula não pode ter mais de 255 caracteres."
                    }, status=400)


            mapa_fatura = {}

            if documento.tipo == 'NC':

                if not documento.documento_origem:
                    return JsonResponse(
                        {"success": False, "error": "Nota de crédito deve ter documento de origem."},
                        status=400
                    )

                artigos_fatura = FinArtigos.objects.filter(id_final=documento.documento_origem, empresa=request.empresa)

                for art in artigos_fatura:
                    mapa_fatura[art.id_art_id] = art

            artigos_data = data.get("artigos", [])
            artigos_validos = []
            if artigos_data:
                for artigo_data in artigos_data:
                    try:
                        artigo_codigo = artigo_data.get("codigo")
                        if not artigo_codigo:
                            logger.error(f"Erro: Código do artigo não fornecido para os dados {artigo_data}")
                            return JsonResponse({"success": False, "error": "Código do artigo não fornecido"},
                                                status=400)

                        artigo_obj = Artigo.objects.filter(id_artigo=artigo_codigo, empresa=request.empresa).first()
                        if not artigo_obj:
                            logger.error(f"Erro: Artigo com código {artigo_codigo} não encontrado.")
                            return JsonResponse(
                                {"success": False, "error": f"Artigo com código {artigo_codigo} não encontrado."},
                                status=400)

                        # Continuar a criação do artigo, mais validações podem ser feitas aqui
                        tipo_artigo = artigo_obj.tipo
                        preco = Decimal(artigo_data.get("preco", "0").replace(",", "."))
                        if preco <= 0:
                            return JsonResponse(
                                {"success": False, "error": "O preço deve ser maior que zero."},
                                status=400
                            )
                        desconto = Decimal(artigo_data.get("desconto", "0").replace(",", "."))
                        if desconto < 0:
                            return JsonResponse(
                                {"success": False, "error": "O desconto não pode ser negativo."},
                                status=400
                            )

                        iva = Decimal(artigo_data.get("iva", "0").replace(",", "."))
                        TAXAS_VALIDAS = {0, 6, 13, 23}

                        if iva not in TAXAS_VALIDAS:
                            return JsonResponse(
                                {"success": False, "error": "Taxa de IVA inválida."},
                                status=400
                            )
                        total = Decimal(artigo_data.get("total", "0").replace(",", "."))


                        quantidade_raw = artigo_data.get("quantidade", 0)

                        try:
                            quantidade = int(quantidade_raw)
                        except (ValueError, TypeError):
                            return JsonResponse(
                                {"success": False, "error": "Quantidade inválida."},
                                status=400
                            )

                        if quantidade < 1 or quantidade > 10000:
                            return JsonResponse(
                                {"success": False, "error": "A quantidade deve estar entre 1 e 10000."},
                                status=400
                            )
                        if desconto > preco * quantidade:
                            return JsonResponse(
                                {"success": False, "error": "O desconto não pode ser superior ao valor da linha."},
                                status=400
                            )
                        descricao_artigo = artigo_data.get("descricao", "").strip()

                        if not descricao_artigo:
                            return JsonResponse(
                                {"success": False, "error": "A descrição do artigo é obrigatória."},
                                status=400
                            )

                        if len(descricao) > 255:
                            return JsonResponse(
                                {"success": False, "error": "A descrição não pode exceder 255 caracteres."},
                                status=400
                            )

                        subtotal = (preco * quantidade - desconto).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)

                        total_recebido = total.quantize(
                            Decimal("0.01"), rounding=ROUND_HALF_UP
                        )

                        if subtotal != total_recebido:
                            return JsonResponse(
                                {
                                    "success": False,
                                    "error": f"Total inválido. Esperado: {subtotal}, Recebido: {total_recebido}"
                                },
                                status=400
                            )

                        motivo_tax = artigo_data.get("motivo")

                        if iva == 0:
                            # Motivo obrigatório
                            if not motivo_tax:
                                return JsonResponse(
                                    {"success": False, "error": "Motivo de IVA obrigatório quando a taxa é 0%."},
                                    status=400
                                )

                            tax_reason = TaxReason.objects.filter(code=motivo_tax).first()
                            if not tax_reason:
                                logger.error(f"Erro: Motivo IVA '{motivo_tax}' não encontrado.")
                                return JsonResponse(
                                    {"success": False, "error": f"Motivo IVA '{motivo_tax}' inválido."},
                                    status=400
                                )

                        else:
                            if motivo_tax:
                                return JsonResponse(
                                    {"success": False, "error": "Motivo de IVA só é permitido quando a taxa é 0%."},
                                    status=400
                                )

                            tax_reason = None

                        if documento.tipo == 'NC':

                            artigo_fatura = mapa_fatura.get(artigo_obj.id_artigo)

                            if not artigo_fatura:
                                return JsonResponse(
                                    {"success": False,
                                     "error": f"O artigo {artigo_codigo} não existe na fatura original."},
                                    status=400
                                )

                            if (
                                    artigo_fatura.preco != preco or
                                    artigo_fatura.desconto != desconto or
                                    artigo_fatura.taxa != iva or
                                    artigo_fatura.descricao != descricao_artigo
                            ):
                                return JsonResponse(
                                    {"success": False,
                                     "error": f"O artigo {artigo_codigo} não corresponde ao da fatura original."},
                                    status=400
                                )

                            if quantidade > artigo_fatura.quantidade:
                                return JsonResponse(
                                    {
                                        "success": False,
                                        "error": f"A quantidade do artigo {artigo_codigo} excede a da fatura original."
                                    },
                                    status=400
                                )

                        artigos_validos.append({
                            "id_art": artigo_obj,
                            "tipo": tipo_artigo,
                            "descricao": descricao_artigo,
                            "quantidade": quantidade,
                            "preco": preco,
                            "desconto": desconto,
                            "taxa": iva,
                            "total": subtotal,
                            "motivo": tax_reason,
                            "empresa": request.empresa
                        })

                    except Exception as e:
                        logger.error(f"Erro ao processar artigo {artigo_data}: {str(e)}")
                        return JsonResponse({"success": False, "error": f"Erro ao processar artigo: {str(e)}"},
                                            status=400)

            artigos_data = data.get("artigos", [])
            if not artigos_data:
                return JsonResponse(
                    {"success": False, "error": "O documento deve conter pelo menos um artigo."},
                    status=400
                )
            valor_total = Decimal("0.00")
            for artigo_data in artigos_data:
                quantidade_raw = artigo_data.get("quantidade", 0)

                try:
                    quantidade = int(quantidade_raw)
                except (ValueError, TypeError):
                    return JsonResponse(
                        {"success": False, "error": "Quantidade inválida."},
                        status=400
                    )

                desconto = Decimal(artigo_data.get("desconto", "0").replace(",", "."))
                preco = Decimal(artigo_data.get("preco", "0").replace(",", "."))
                iva = Decimal(artigo_data.get("iva", "0").replace(",", "."))
                subtotal = (preco * quantidade - desconto).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
                total_com_iva = (subtotal * (1 + iva / 100)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
                valor_total += total_com_iva


            if documento.tipo == 'NC':
                    valor_total = -abs(valor_total)

            if documento.tipo in ['FR', 'FS']:
                total_pago = valor_total
            else:
                total_pago = 0
            if documento.tipo == 'FS':
                limite_consumidor_final = Decimal("1000.00")
                limite_sujeito_passivo = Decimal("100.00")

                tem_nif = False
                if documento.cliente and documento.cliente.contribuinte:
                    if str(documento.cliente.contribuinte) != "999999990":
                        tem_nif = True

                if tem_nif and valor_total > limite_sujeito_passivo:
                    return JsonResponse({
                        "success": False,
                        "error": f"Fatura Simplificada para empresas não pode exceder {limite_sujeito_passivo}€. Use Fatura (FT)."
                    }, status=400)

                if valor_total > limite_consumidor_final:
                    return JsonResponse({
                        "success": False,
                        "error": f"Fatura Simplificada não pode exceder {limite_consumidor_final}€. Use Fatura (FT)."
                    }, status=400)

            ano = documento.data_emissao.year
            serie = documento.serie
            tipo = documento.tipo

            numero_final = obter_proximo_numero_final(tipo, serie, ano, request.empresa)

            empr = Empresa.objects.first()
            codigo_at_tributaria = gerar_codigo_at(documento)
            documento.valor_total = valor_total
            documento_final = DocumentoFinalizado.objects.create(
                tipo=tipo,
                serie=serie,
                numero=numero_final,
                ano=ano,

                cliente_id=documento.cliente.id_cliente if documento.cliente else None,
                cliente_nome=documento.cliente.nome if documento.cliente else "",
                cliente_morada1=documento.cliente.morada1 if documento.cliente else "",
                cliente_morada2=documento.cliente.morada2 if documento.cliente else "",
                cliente_codigo_postal=documento.cliente.codigo_postal if documento.cliente else "",
                cliente_concelho=documento.cliente.concelho if documento.cliente else "",
                cliente_pais=documento.cliente.pais if documento.cliente else "",
                cliente_email=documento.cliente.email if documento.cliente else "",
                cliente_contribuinte=documento.cliente.contribuinte if documento.cliente else "",

                empresa_nome=empr.nome,
                empresa_morada=empr.morada,
                empresa_codigo_postal=empr.codigo_postal,
                empresa_cidade=empr.cidade,
                empresa_pais=empr.pais,
                empresa_email=empr.email,
                empresa_contribuinte=empr.nif,

                modalidade_nome=pagamento_obj,
                moeda_simbolo=moeda_obj.simbolo,
                ordem_compra=ordem_compra,
                numero_compromisso=numero_compromisso,
                descricao=documento.descricao,
                rodape=rodape,
                data_emissao=documento.data_emissao,
                data_vencimento=documento.data_vencimento,
                valor_total=valor_total,
                total_pago=total_pago,
                local_carga=local_carga,
                local_descarga=local_descarga,
                data_carga=documento.data_carga,
                data_descarga=documento.data_descarga,
                expedicao=expedicao,
                transporte_descricao=transporte_descricao,
                codigo_at_tributaria=codigo_at_tributaria,
                documento_origem=documento.documento_origem if documento.documento_origem else None,
                empresa=request.empresa,

                estado='Finalizado'
            )

            if documento_final.tipo in ['FR', 'FS']:
                try:
                    ano_atual = timezone.now().year
                    ultimo_re = Recibo.objects.filter(ano=ano_atual, empresa=request.empresa).order_by('-numero').first()
                    novo_num_re = (ultimo_re.numero + 1) if ultimo_re else 1

                    novo_recibo = Recibo.objects.create(
                        tipo='RE',
                        serie=str(ano_atual),
                        numero=novo_num_re,
                        ano=ano_atual,
                        cliente_id=documento_final.cliente_id,
                        data_emissao=documento_final.data_emissao,
                        modalidade_nome=documento_final.modalidade_nome.nome if documento_final.modalidade_nome else "Numerário",
                        valor_total=documento_final.valor_total,
                        estado='Normal',
                        empresa=request.empresa
                    )

                    ReciboLinhas.objects.create(
                        id_recibo_final=novo_recibo,
                        id_doc_final=documento_final,
                        documento_tipo=documento_final.tipo,
                        documento_numero=f"{documento_final.serie}-{documento_final.numero}/{documento_final.ano}",
                        data_emissao=documento_final.data_emissao,
                        valor_documento=documento_final.valor_total,
                        valor_recebido=documento_final.valor_total,
                        valor_em_divida=0,
                        empresa=request.empresa
                    )

                    logger.info(f"Recibo {novo_recibo.numero} gerado automaticamente para {documento_final.tipo}")

                except Exception as re_err:
                    logger.error(f"Erro ao gerar recibo automático: {str(re_err)}")

            for artigo in artigos_validos:
                FinArtigos.objects.create(
                    id_final=documento_final,
                    id_art=artigo["id_art"],
                    tipo=artigo["tipo"],
                    descricao=artigo["descricao"],
                    quantidade=artigo["quantidade"],
                    preco=artigo["preco"],
                    desconto=artigo["desconto"],
                    taxa=artigo["taxa"],
                    total=artigo["total"],
                    motivo_id=artigo["motivo"].id if artigo["motivo"] else None,
                    empresa=artigo["empresa"]
                )

            documento.delete()

            return JsonResponse({
                "success": True,
                "id": documento_final.id,
                "message": "Documento finalizado com sucesso!"
            })

        except ValidationError as e:
            return JsonResponse(
                {"success": False, "error": "; ".join(e.messages)},
                status=400
            )
        except Exception as e:
            print(f"ERRO NO BACKEND: {str(e)}")  # Isto vai aparecer no seu terminal preto
            import traceback
            traceback.print_exc()  # Isto mostra a linha exata onde falhou
            return JsonResponse({"success": False, "error": str(e)}, status=400)
    else:
        return JsonResponse(
            {"success": False, "error": "Método não permitido. Use POST."},
            status=405
        )


@login_required
@empresa_obrigatoria
def ver_fatura(request, id):
    documento = get_object_or_404(DocumentoFinalizado, id=id, empresa=request.empresa)

    nota_credito = DocumentoFinalizado.objects.filter(documento_origem=documento, empresa=request.empresa).exclude(estado='Anulado').first()

    if nota_credito:
        valor_nota_credito = nota_credito.valor_total
    else:
        valor_nota_credito = 0

    return render(
        request,
        "faturas/ver_faturas.html",
        {
            'doc_id': id,
            'id_cliente': documento.cliente_id,
            'tipo_documento': documento.tipo,
            'valor_em_divida': documento.valor_total - documento.total_pago,
            'modalidades': Modalidade.objects.all(),
            'valor_nota_credito': valor_nota_credito,
            'moeda_tipo': documento.moeda_simbolo
        }
    )

@login_required
@empresa_obrigatoria
def api_documento_completo(request, doc_id):
    doc = get_object_or_404(DocumentoFinalizado, id=doc_id, empresa=request.empresa)
    artigos = FinArtigos.objects.filter(id_final=doc, empresa=request.empresa)

    lista_artigos = [{
        "tipo": a.tipo,
        "id": a.id_art.id_artigo,
        "descricao": a.descricao,
        "quantidade": a.quantidade,
        "preco": float(a.preco),
        "desconto": float(a.desconto),
        "taxa": float(a.taxa),
        "motivo": a.motivo.code if a.motivo else ""
    } for a in artigos]

    data = {
        "id": doc.id,
        "tipo": doc.tipo,
        "serie": doc.serie,
        "ano": doc.ano,
        "numero": doc.numero,
        "data_emissao": doc.data_emissao.isoformat(),
        "data_vencimento": doc.data_vencimento.isoformat(),
        "descricao": doc.descricao,
        "rodape": doc.rodape,
        "ordem_compra": doc.ordem_compra,
        "numero_compromisso": doc.numero_compromisso,
        "local_carga": doc.local_carga,
        "local_descarga": doc.local_descarga,
        "data_carga": doc.data_carga,
        "data_descarga": doc.data_descarga,
        "expedicao": doc.expedicao,
        "matricula": doc.transporte_descricao,
        "total_pago": doc.total_pago,
        "total_fatura":doc.valor_total,
        "empresa": {
            "nome": doc.empresa_nome,
            "morada": doc.empresa_morada,
            "codigo_postal": doc.empresa_codigo_postal,
            "cidade": doc.empresa_cidade,
            "pais": doc.empresa_pais,
            "email": doc.empresa_email,
            "nif": doc.empresa_contribuinte
        },

        "cliente": {
            "nome": doc.cliente_nome,
            "morada1": doc.cliente_morada1,
            "morada2": doc.cliente_morada2,
            "codigo_postal": doc.cliente_codigo_postal,
            "concelho": doc.cliente_concelho,
            "pais": doc.cliente_pais,
            "email": doc.cliente_email,
            "contribuinte": doc.cliente_contribuinte,
            "modalidade": {"nome": doc.modalidade_nome} if doc.modalidade_nome else None,
        },

        "moeda": {"id": 1, "simbolo": doc.moeda_simbolo},  # Ajuste o ID conforme sua lógica
        "moedas": list(Moeda.objects.values('id', 'simbolo')),  # Para preencher o select
        "modalidades": list(Modalidade.objects.values('id_modalidade', 'nome')),  # Para o select de pagamento
        "tax_reasons": list(TaxReason.objects.values('code', 'description')),  # Para o select de isenção
        "artigos": lista_artigos
    }
    return JsonResponse(data)

from django.shortcuts import render

from django.db.models import Sum
from .models import Cliente1, Vendedor, Zona, Transporte, Impostos, Pagamento, Modalidade, Precos

@login_required
@empresa_obrigatoria
def cliente_detalhes(request, id_cliente):
    cliente = get_object_or_404(Cliente1, id_cliente=id_cliente, empresa=request.empresa)

    todos_docs = DocumentoFinalizado.objects.filter(cliente_id=id_cliente, empresa=request.empresa)

    faturas_faturadas = todos_docs.exclude(tipo__in=['GT', 'NC'])
    faturas_financeiras = todos_docs.exclude(tipo__in=['GT'])
    guias_transporte = todos_docs.filter(tipo='GT')

    recibos = Recibo.objects.filter(cliente_id=id_cliente, empresa=request.empresa).order_by('-criado_em')

    total_faturado = faturas_faturadas.aggregate(Sum('valor_total'))['valor_total__sum'] or 0

    total_nota_credito = 0
    for fatura in faturas_financeiras:
        notas_credito = DocumentoFinalizado.objects.filter(documento_origem=fatura, empresa=request.empresa).exclude(estado='Anulado')

        total_nota_credito += sum(nota_credito.valor_total for nota_credito in notas_credito)

    total_faturado_com_desconto = total_faturado + total_nota_credito

    total_via_recibos = Recibo.objects.filter(
        cliente_id=id_cliente,
        estado='Normal',
        empresa=request.empresa
    ).aggregate(Sum('valor_total'))['valor_total__sum'] or 0

    total_recebido = total_via_recibos

    saldo_pendente = total_faturado_com_desconto - total_recebido

    context = {
        "cliente": cliente,
        "faturas_financeiras": faturas_financeiras,
        "guias_transporte": guias_transporte,
        "total_faturado": total_faturado,
        'total_recebido': total_recebido,
        'saldo_pendente': saldo_pendente,
        "vendedores": Vendedor.objects.all(),
        "zonas": Zona.objects.all(),
        "transportes": Transporte.objects.filter(empresa=request.empresa),
        "impostos_list": Impostos.objects.all(),
        "pagamentos": Pagamento.objects.all(),
        "modalidades": Modalidade.objects.all(),
        "precos_list": Precos.objects.all(),
        "recibos": recibos,
    }

    return render(request, "subsubconteudo/cliente_detalhes.html", context)

import csv
@login_required
@empresa_obrigatoria
def exportar_cliente_csv(request, id_cliente):
    cliente = Cliente1.objects.get(id_cliente=id_cliente, empresa=request.empresa)
    facturas = DocumentoFinalizado.objects.filter(cliente_id=id_cliente, empresa=request.empresa)

    NOMES_TIPOS = {
        'FT': 'Fatura',
        'FR': 'Fatura Recibo',
        'FS': 'Fatura Simplificada',
        'NC': 'Nota de Crédito',
    }
    # 2. Configurar o Response para download de CSV
    response = HttpResponse(content_type='text/csv; charset=utf-8-sig')  # utf-8-sig para o Excel abrir bem os acentos
    response['Content-Disposition'] = f'attachment; filename="dados_cliente_{cliente.contribuinte}.csv"'

    writer = csv.writer(response, delimiter=';')  # Ponto e vírgula costuma ser melhor para o Excel em PT

    # 3. ESCREVER DADOS DO CLIENTE (Cabeçalho + Dados)
    writer.writerow(['--- DADOS DO CLIENTE ---'])
    writer.writerow(['NIF', 'Nome', 'Email', 'Morada', 'CP', 'Cidade', 'Pais'])
    writer.writerow([
        cliente.contribuinte,
        cliente.nome,
        cliente.email,
        cliente.morada1,
        cliente.codigo_postal,
        cliente.concelho,
        cliente.pais
    ])

    writer.writerow([])  # Linha em branco para separar

    # 4. ESCREVER HISTÓRICO DE FATURAS
    writer.writerow(['--- HISTÓRICO DE FATURAS ---'])
    writer.writerow(['Tipo', 'Numero', 'Data Emissao', 'Valor Total', 'Moeda', 'Pais na Fatura', 'Estado'])

    for f in facturas:
        # 1. Converter a sigla para o nome por extenso
        tipo_extenso = NOMES_TIPOS.get(f.tipo, f.tipo)



        writer.writerow([
            tipo_extenso,
            f"{f.serie}-{f.numero}/{f.ano}",
            f.data_emissao,
            str(f.valor_total).replace('.', ','),  # Trocar ponto por vírgula para o Excel PT
            f.moeda_simbolo,
            f.cliente_pais,
            f.estado
        ])

    return response

@login_required
@empresa_obrigatoria
def guias_json(request):
    TIPO_EXTENSO = {
        'GT': 'Guia de Transporte',
    }
    documentos = (
        DocumentoFinalizado.objects
        .filter(tipo__in=TIPO_EXTENSO.keys(),
                empresa=request.empresa)
        .order_by('-id', '-data_emissao')
    )

    data = []
    for doc in documentos:
        numero_completo = f"{doc.tipo}/{doc.serie}/{doc.ano}/{doc.numero}"

        data.append({
            "id_documento": doc.id,
            "tipo": TIPO_EXTENSO.get(doc.tipo, doc.tipo),
            "numero": numero_completo,
            "cliente_nome": doc.cliente_nome,
            "data_emissao": doc.data_emissao.strftime('%Y-%m-%d') if doc.data_emissao else "",
            "local_descarga": doc.local_descarga or "---",
            "temporario": False
        })

    return JsonResponse(data, safe=False)

@login_required
@empresa_obrigatoria
@csrf_exempt
def reservar_numero_guia(request):
    if request.method != "POST":
        return JsonResponse({"erro": "Método inválido."}, status=405)

    cliente_id = request.POST.get("cliente")
    documento = request.POST.get("documento", "GT") # Default para Guia de Transporte
    data_emissao = request.POST.get("data_emissao")

    if not all([cliente_id, data_emissao]):
        return JsonResponse({"erro": "Cliente e Data são obrigatórios."}, status=400)

    ano = int(data_emissao[:4])

    try:
        with transaction.atomic():
            contador, created = DocumentoContador.objects.select_for_update().get_or_create(
                empresa=request.empresa,
                tipo=documento,
                ano=ano,
                defaults={"serie": "G", "ultimo_numero": 0}
            )

            proximo_numero = contador.ultimo_numero + 1
            contador.ultimo_numero = proximo_numero
            contador.save()

        return JsonResponse({
            "success": True,
            "cliente_id": cliente_id,
            "tipo": documento,
            "serie": "G",
            "numero": proximo_numero,
            "ano": ano,
            "data_emissao": data_emissao
        })
    except Exception as e:
        return JsonResponse({"erro": str(e)}, status=500)

@login_required
@empresa_obrigatoria
def guia_documento(request):
    cliente_id = request.GET.get('cliente')
    tipo = request.GET.get('tipo', 'GT')
    serie = request.GET.get('serie')
    numero = request.GET.get('numero')
    ano = request.GET.get('ano')
    data_emissao = request.GET.get('data')

    # Busca o cliente para mostrar o nome/NIF no topo da página
    cliente = get_object_or_404(Cliente1, id_cliente=cliente_id, empresa=request.empresa)

    context = {
        'cliente': cliente,
        'tipo': tipo,
        'serie': serie,
        'numero': numero,
        'ano': ano,
        'data_emissao': data_emissao,
        'numero_completo': f"{tipo}/{serie}/{ano}/{numero}"
    }

    return render(request, 'guias/nova_guia.html', context)

@login_required
@empresa_obrigatoria
def api_guia_preparar(request):
    cliente_id = request.GET.get('cliente')
    moeda_id = request.GET.get('moeda', 1)

    cliente_data = None
    if cliente_id:
        c = get_object_or_404(Cliente1.objects.select_related("transporte", "modalidade", "impostos"),
                              id_cliente=cliente_id,
                              empresa=request.empresa)
        cliente_data = {
            "id": c.id_cliente,
            "codigo": c.codigo,
            "nome": c.nome,
            "morada1": c.morada1,
            "morada2": c.morada2,
            "codigo_postal": c.codigo_postal,
            "pais": c.pais,
            "concelho": c.concelho,
            "contribuinte": c.contribuinte,
            "email": c.email,
            "transporte": c.transporte.descricao if c.transporte else "",
            "impostos": c.impostos.nome if c.impostos else "IVA",
            "local_descarga": c.morada1,
        }

    # 4. Dados de Apoio (Selects)
    moedas = list(Moeda.objects.values("id", "codigo", "nome", "simbolo"))
    tax_reasons = list(TaxReason.objects.values("code", "description"))

    moeda_sel = Moeda.objects.filter(id=moeda_id).first() or Moeda.objects.first()

    empresa=request.empresa
    return JsonResponse({
        # --- DOCUMENTO (Rascunho vindo da URL) ---
        "id": None,  # Não existe ID temp ainda
        "tipo": request.GET.get('tipo', 'GT'),
        "serie": request.GET.get('serie', 'G'),
        "numero": request.GET.get('numero'),
        "ano": request.GET.get('ano'),
        "ordem_compra": "",
        "numero_compromisso": "",
        "descricao": "",
        "rodape": "",
        "data_emissao": request.GET.get('data'),
        "valor_total": 0.00,

        "expedicao": "Viatura Própria",
        "matricula": "",

        "cliente": cliente_data,

        "moeda": {
            "id": moeda_sel.id,
            "codigo": moeda_sel.codigo,
            "nome": moeda_sel.nome,
            "simbolo": moeda_sel.simbolo,
        },
        "moedas": moedas,

        "empresa": {
            "nome": empresa.nome,
            "morada": empresa.morada,
            "codigo_postal": empresa.codigo_postal,
            "cidade": empresa.cidade,
            "pais": empresa.pais,
            "email": empresa.email,
            "nif": empresa.nif,
            "local": empresa.local
        },

        "tax_reasons": tax_reasons,
        "artigos": []
    })

@login_required
@empresa_obrigatoria
def finalizar_documento_guia(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            ordem_compra = data.get("ordem_compra")
            local_carga = data.get("local_carga")
            local_descarga = data.get("local_descarga")
            tipo = data.get("tipo")

            if ordem_compra:
                ordem_compra = validar_ordem_compra(ordem_compra)

            numero_compromisso = data.get("numero_comp")
            if numero_compromisso:
               numero_compromisso = validar_numero_compromisso(numero_compromisso)

            descricao = data.get("descricao")
            if descricao:
                descricao = validar_texto_longo(descricao, limite=200)

            rodape = data.get("rodape")
            if rodape:
                rodape = validar_texto_longo(rodape, limite=200)
            try:
                data_emissao = data.get("data_emissao")
                if data_emissao:
                    data_emissao = datetime.strptime(data_emissao, "%d/%m/%Y")

                    # Verificar se a data de emissão é maior ou igual à data atual
                    if data_emissao.date() < timezone.now().date():
                        return JsonResponse(
                            {"success": False, "error": "A data de emissão não pode ser anterior à data atual."},
                            status=400)

                if tipo in ("GT"):
                    if len(local_carga) > 255:
                        raise ValidationError("O local de carga não pode ter mais que 255 caracteres.")

                    if len(local_descarga) > 255:
                        raise ValidationError("O local de descarga não pode ter mais que 255 caracteres.")

                data_carga_str = data.get("data_carga")
                data_descarga_str = data.get("data_descarga")

                data_carga = ""
                data_descarga = ""

                if data_carga_str:
                    dt_carga_naive = datetime.strptime(data_carga_str, "%d/%m/%Y %H:%M")
                    data_carga = timezone.make_aware(
                        dt_carga_naive,
                        timezone.get_current_timezone()
                    )

                    if data_carga.date() < timezone.now().date():
                        return JsonResponse(
                            {"success": False, "error": "A data de carga não pode ser anterior à data atual."},
                            status=400
                        )

                if data_descarga_str:
                    dt_descarga_naive = datetime.strptime(data_descarga_str, "%d/%m/%Y %H:%M")
                    data_descarga = timezone.make_aware(
                        dt_descarga_naive,
                        timezone.get_current_timezone()
                    )

                    if data_descarga < timezone.now():
                        return JsonResponse(
                            {"success": False, "error": "A data de descarga não pode ser anterior à data atual."},
                            status=400
                        )

                    if data_carga_str and data_descarga < data_carga:
                        return JsonResponse(
                            {"success": False, "error": "A data de descarga não pode ser anterior à data de carga."},
                            status=400
                        )

                    data_descarga = data_descarga
            except ValueError as e:
                return JsonResponse({"success": False, "error": f"Erro no formato da data: {str(e)}"}, status=400)

            if data_carga and data_emissao:
                if data_carga.date() < data_emissao.date():
                    return JsonResponse(
                        {"success": False, "error": "A data de carga não pode ser anterior à data de emissão."},
                        status=400
                    )
            expedicao = data.get("expedicao")
            if expedicao:
                if len(expedicao) > 255:
                    raise ValidationError("O modo de expedição não pode ter mais que 255 caracteres.")

            cliente_id = data.get("cliente_id")
            if not cliente_id:
                return JsonResponse(
                    {"success": False, "error": "Cliente obrigatório."},
                    status=400
                )
            if cliente_id:
                cliente_obj = Cliente1.objects.filter(id_cliente=cliente_id, empresa=request.empresa).first()
                if not cliente_obj:
                    return JsonResponse({"success": False, "error": f"Cliente {cliente_id} não encontrado."}, status=400)
                cliente_id = cliente_obj.id_cliente

            moeda_obj = None
            moeda_id = data.get("moeda")
            if moeda_id:
                moeda_obj = Moeda.objects.filter(id=moeda_id).first()
                if not moeda_obj:
                    return JsonResponse({"success": False, "error": f"Moeda {moeda_id} não encontrada."}, status=400)
                moeda_id = moeda_obj.id

            transporte_descricao = data.get("matricula")


            if transporte_descricao:
                transporte_obj = Transporte.objects.filter(
                    descricao=transporte_descricao,
                    empresa=request.empresa
                ).first()

                if not transporte_obj:
                    return JsonResponse(
                        {"success": False,
                         "error": f"Transporte com matrícula {transporte_descricao} não encontrado."},
                        status=400
                    )


            campos_obrigatorios = [
                'local_carga', 'local_descarga',
                'data_carga', 'data_descarga', 'matricula'
            ]
            for campo in campos_obrigatorios:
                valor = data.get(campo)
                if not valor or str(valor).strip() == "":
                    return JsonResponse({
                        "error": f"O campo {campo.replace('_', ' ')} é obrigatório."
                    }, status=400)

            artigos_data = data.get("artigos", [])
            artigos_validos = []
            if artigos_data:
                for artigo_data in artigos_data:
                    try:
                        # Aqui pode ocorrer um erro se 'codigo' não for fornecido ou o formato de dados estiver errado
                        artigo_codigo = artigo_data.get("codigo")
                        if not artigo_codigo:
                            logger.error(f"Erro: Código do artigo não fornecido para os dados {artigo_data}")
                            return JsonResponse({"success": False, "error": "Código do artigo não fornecido"},
                                                status=400)

                        artigo_obj = Artigo.objects.filter(id_artigo=artigo_codigo, empresa=request.empresa).first()
                        if not artigo_obj:
                            logger.error(f"Erro: Artigo com código {artigo_codigo} não encontrado.")
                            return JsonResponse(
                                {"success": False, "error": f"Artigo com código {artigo_codigo} não encontrado."},
                                status=400)

                        # Continuar a criação do artigo, mais validações podem ser feitas aqui
                        tipo_artigo = artigo_obj.tipo
                        preco = Decimal(artigo_data.get("preco", "0").replace(",", "."))
                        if preco < 0:
                            return JsonResponse(
                                {"success": False, "error": "O preço deve ser maior que zero."},
                                status=400
                            )
                        desconto = Decimal(artigo_data.get("desconto", "0").replace(",", "."))
                        if desconto < 0:
                            return JsonResponse(
                                {"success": False, "error": "O desconto não pode ser negativo."},
                                status=400
                            )

                        iva = Decimal(artigo_data.get("iva", "0").replace(",", "."))
                        TAXAS_VALIDAS = {0, 6, 13, 23}

                        if iva not in TAXAS_VALIDAS:
                            return JsonResponse(
                                {"success": False, "error": "Taxa de IVA inválida."},
                                status=400
                            )
                        total = Decimal(artigo_data.get("total", "0").replace(",", "."))


                        quantidade_raw = artigo_data.get("quantidade", 0)

                        try:
                            quantidade = int(quantidade_raw)
                        except (ValueError, TypeError):
                            return JsonResponse(
                                {"success": False, "error": "Quantidade inválida."},
                                status=400
                            )

                        if quantidade < 1 or quantidade > 10000:
                            return JsonResponse(
                                {"success": False, "error": "A quantidade deve estar entre 1 e 10000."},
                                status=400
                            )
                        if desconto > preco * quantidade:
                            return JsonResponse(
                                {"success": False, "error": "O desconto não pode ser superior ao valor da linha."},
                                status=400
                            )
                        descricao_artigo = artigo_data.get("descricao", "").strip()

                        if not descricao_artigo:
                            return JsonResponse(
                                {"success": False, "error": "A descrição do artigo é obrigatória."},
                                status=400
                            )

                        if len(descricao) > 255:
                            return JsonResponse(
                                {"success": False, "error": "A descrição não pode exceder 255 caracteres."},
                                status=400
                            )

                        subtotal = (preco * quantidade - desconto).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)

                        total_recebido = total.quantize(
                            Decimal("0.01"), rounding=ROUND_HALF_UP
                        )

                        if subtotal != total_recebido:
                            return JsonResponse(
                                {
                                    "success": False,
                                    "error": f"Total inválido. Esperado: {subtotal}, Recebido: {total_recebido}"
                                },
                                status=400
                            )

                        motivo_tax = artigo_data.get("motivo")

                        if iva == 0:
                            # Motivo obrigatório
                            if not motivo_tax:
                                return JsonResponse(
                                    {"success": False, "error": "Motivo de IVA obrigatório quando a taxa é 0%."},
                                    status=400
                                )

                            tax_reason = TaxReason.objects.filter(code=motivo_tax).first()
                            if not tax_reason:
                                logger.error(f"Erro: Motivo IVA '{motivo_tax}' não encontrado.")
                                return JsonResponse(
                                    {"success": False, "error": f"Motivo IVA '{motivo_tax}' inválido."},
                                    status=400
                                )

                        else:
                            if motivo_tax:
                                return JsonResponse(
                                    {"success": False, "error": "Motivo de IVA só é permitido quando a taxa é 0%."},
                                    status=400
                                )

                            tax_reason = None

                        artigos_validos.append({
                            "id_art": artigo_obj,
                            "tipo": tipo_artigo,
                            "descricao": descricao_artigo,
                            "quantidade": quantidade,
                            "preco": preco,
                            "desconto": desconto,
                            "taxa": iva,
                            "total": subtotal,
                            "motivo": tax_reason,
                        })

                    except Exception as e:
                        logger.error(f"Erro ao processar artigo {artigo_data}: {str(e)}")
                        return JsonResponse({"success": False, "error": f"Erro ao processar artigo: {str(e)}"},
                                            status=400)

            artigos_data = data.get("artigos", [])
            if not artigos_data:
                return JsonResponse(
                    {"success": False, "error": "O documento deve conter pelo menos um artigo."},
                    status=400
                )
            valor_total = Decimal("0.00")
            for artigo_data in artigos_data:
                quantidade_raw = artigo_data.get("quantidade", 0)

                try:
                    quantidade = int(quantidade_raw)
                except (ValueError, TypeError):
                    return JsonResponse(
                        {"success": False, "error": "Quantidade inválida."},
                        status=400
                    )

                desconto = Decimal(artigo_data.get("desconto", "0").replace(",", "."))
                preco = Decimal(artigo_data.get("preco", "0").replace(",", "."))
                iva = Decimal(artigo_data.get("iva", "0").replace(",", "."))
                subtotal = (preco * quantidade - desconto).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
                total_com_iva = (subtotal * (1 + iva / 100)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
                valor_total += total_com_iva


            id_cliente = data.get("cliente_id")
            cliente = Cliente1.objects.filter(id_cliente=id_cliente, empresa=request.empresa).first()

            if not cliente:
                raise ValueError("Cliente não encontrado")

            serie = data.get("serie")
            numero_raw = data.get("numero")

            if not numero_raw:
                return JsonResponse(
                    {"success": False, "error": "Número do documento é obrigatório."},
                    status=400
                )

            try:
                numero = int(numero_raw)
            except (ValueError, TypeError):
                return JsonResponse(
                    {"success": False, "error": "Número do documento inválido."},
                    status=400
                )

            ano = data_emissao.year
            serie = data.get("serie")

            numero_final = obter_proximo_numero_final(tipo, serie, ano, request.empresa)

            empr = request.empresa
            with transaction.atomic():
                documento_final = DocumentoFinalizado.objects.create(
                    tipo=tipo,
                    serie=serie,
                    numero=numero_final,
                    ano=ano,

                    cliente_id=id_cliente,
                    cliente_nome=cliente.nome if cliente else "",
                    cliente_morada1=cliente.morada1 if cliente else "",
                    cliente_morada2=cliente.morada2 if cliente else "",
                    cliente_codigo_postal=cliente.codigo_postal if cliente else "",
                    cliente_concelho=cliente.concelho if cliente else "",
                    cliente_pais=cliente.pais if cliente else "",
                    cliente_email=cliente.email if cliente else "",
                    cliente_contribuinte=cliente.contribuinte if cliente else "",

                    empresa_nome=empr.nome,
                    empresa_morada=empr.morada,
                    empresa_codigo_postal=empr.codigo_postal,
                    empresa_cidade=empr.cidade,
                    empresa_pais=empr.pais,
                    empresa_email=empr.email,
                    empresa_contribuinte=empr.nif,

                    moeda_simbolo=moeda_obj.simbolo,
                    ordem_compra=ordem_compra,
                    numero_compromisso=numero_compromisso,
                    descricao=descricao,
                    rodape=rodape,
                    data_emissao=data_emissao,
                    data_vencimento=data_emissao + timedelta(days=30),
                    valor_total=valor_total,
                    total_pago=0,
                    local_carga=local_carga,
                    local_descarga=local_descarga,
                    data_carga=data_carga,
                    data_descarga=data_descarga,
                    expedicao=expedicao,
                    transporte_descricao=transporte_descricao,
                    codigo_at_tributaria= "",
                    empresa=request.empresa,

                    estado='Finalizado'
                )

                for artigo in artigos_validos:
                    FinArtigos.objects.create(
                        id_final=documento_final,
                        id_art=artigo["id_art"],
                        tipo=artigo["tipo"],
                        descricao=artigo["descricao"],
                        quantidade=artigo["quantidade"],
                        preco=artigo["preco"],
                        desconto=artigo["desconto"],
                        taxa=artigo["taxa"],
                        total=artigo["total"],
                        motivo_id=artigo["motivo"].id if artigo["motivo"] else None,
                        empresa=request.empresa
                    )

                codigo_at = gerar_codigo_at(documento_final)
                documento_final.codigo_at_tributaria = codigo_at
                documento_final.save(update_fields=["codigo_at_tributaria"])

            return JsonResponse({
                "success": True,
                "id": documento_final.id,
                "message": "Documento finalizado com sucesso!"
            })

        except ValidationError as e:
            return JsonResponse(
                {"success": False, "error": "; ".join(e.messages)},
                status=400
            )
        except Exception as e:
            print(f"ERRO NO BACKEND: {str(e)}")  # Isto vai aparecer no seu terminal preto
            import traceback
            traceback.print_exc()  # Isto mostra a linha exata onde falhou
            return JsonResponse({"success": False, "error": str(e)}, status=400)
    else:
        return JsonResponse(
            {"success": False, "error": "Método não permitido. Use POST."},
            status=405
        )
from django.db.models import F

@login_required
@empresa_obrigatoria
@transaction.atomic
def criar_recibo_cliente(request, id):
    if request.method != "POST":
        return JsonResponse({"error": "Método não permitido. Utilize POST."}, status=405)

    metodo_nome = request.POST.get('metodo')

    # 1. Verifica se o campo foi preenchido
    if not metodo_nome:
        return JsonResponse({"error": "O método de pagamento é obrigatório."}, status=400)

    # 2. Verifica se o método existe na base de dados
    modalidade_valida = Modalidade.objects.filter(nome=metodo_nome).exists()
    if not modalidade_valida:
        return JsonResponse({"error": "Método de pagamento inválido ou não autorizado."}, status=400)

    try:
        valor_str = request.POST.get('valor', '0').replace(',', '.')
        valor_disponivel = Decimal(valor_str)

        if valor_disponivel <= 0:
            return JsonResponse({"error": "O valor do recibo deve ser superior a zero."}, status=400)

        # Buscar faturas pendentes
        faturas_pendentes = DocumentoFinalizado.objects.filter(
            cliente_id=id,
            estado='Finalizado',
            empresa=request.empresa
        ).exclude(tipo__in=['GT', 'FR']).filter(
            total_pago__lt=F('valor_total')
        ).order_by('data_emissao', 'id')

        if not faturas_pendentes.exists():
            return JsonResponse({"error": "Este cliente não possui faturas pendentes de pagamento."}, status=400)

        # Calcular o total da dívida restante antes de criar o recibo
        total_divida_restante = 0
        for fatura in faturas_pendentes:
            # Buscar notas de crédito para a fatura
            notas_credito = DocumentoFinalizado.objects.filter(documento_origem=fatura, empresa=request.empresa).exclude(estado='Anulado')

            # Calcular o valor total das notas de crédito
            valor_nota_credito = sum(nota_credito.valor_total for nota_credito in notas_credito)

            if valor_nota_credito > 0:
                valor_nota_credito = 0

            # Calcular a dívida atual, considerando as notas de crédito
            divida_atual = fatura.valor_total - fatura.total_pago + valor_nota_credito

            # Garantir que a dívida não seja negativa
            if divida_atual < 0:
                divida_atual = 0

            # Atualizar o total de dívida restante
            total_divida_restante += divida_atual

        # Verificar se o valor do recibo não excede a dívida total
        if valor_disponivel > total_divida_restante:
            return JsonResponse(
                {"error": f"O valor do recibo ({valor_disponivel}) excede a dívida total ({total_divida_restante})."},
                status=400)

        # Gerar numeração sequencial do recibo
        ano_atual = timezone.now().year
        ultimo = Recibo.objects.filter(ano=ano_atual, empresa=request.empresa).order_by('-numero').first()
        novo_numero = (ultimo.numero + 1) if ultimo else 1

        # Criar o cabeçalho do recibo
        recibo = Recibo.objects.create(
            tipo='RE',
            serie=str(ano_atual),
            numero=novo_numero,
            ano=ano_atual,
            cliente_id=id,
            data_emissao=timezone.now().date(),
            modalidade_nome=metodo_nome,
            valor_total=valor_disponivel,
            estado='Normal',
            empresa=request.empresa
        )

        temp_valor = valor_disponivel  # Valor total disponível para abatimento

        for fatura in faturas_pendentes:
            if temp_valor <= 0:
                break

            # Buscar notas de crédito para a fatura
            notas_credito = DocumentoFinalizado.objects.filter(documento_origem=fatura, empresa=request.empresa).exclude(estado='Anulado')

            # Calcular o valor total das notas de crédito
            valor_nota_credito = sum(nota_credito.valor_total for nota_credito in notas_credito)

            # Calcular a dívida atual, considerando as notas de crédito
            divida_atual = fatura.valor_total - fatura.total_pago + valor_nota_credito

            # Garantir que a dívida não seja negativa
            if divida_atual < 0:
                divida_atual = 0

            # Calcular o valor a ser abatido
            valor_a_abater = min(temp_valor, divida_atual)

            # Criar a linha do recibo
            ReciboLinhas.objects.create(
                id_recibo_final=recibo,
                id_doc_final=fatura,
                documento_tipo=fatura.tipo,
                documento_numero=f"{fatura.serie}-{fatura.numero}/{fatura.ano}",
                data_emissao=fatura.data_emissao,
                valor_documento=fatura.valor_total,
                valor_recebido=valor_a_abater,
                valor_em_divida=divida_atual - valor_a_abater,
                empresa=request.empresa
            )

            # Atualizar o total pago da fatura
            fatura.total_pago += valor_a_abater
            fatura.save()

            # Subtrair o valor abatido do valor disponível para pagamento
            temp_valor -= valor_a_abater

        # Retorno com sucesso
        return JsonResponse({
            "success": True,
            "message": "Recibo emitido e faturas liquidadas.",
            "recibo_id": recibo.id
        })

    except IntegrityError:
        return JsonResponse({"error": "Erro de concorrência na numeração. Tente novamente."}, status=500)
    except Exception as e:
        return JsonResponse({"error": f"Erro interno: {str(e)}"}, status=500)


from decimal import Decimal
from django.db import IntegrityError, transaction
from django.http import JsonResponse
from .models import DocumentoFinalizado, Recibo, ReciboLinhas, Modalidade
from django.utils import timezone

@login_required
@empresa_obrigatoria
@transaction.atomic
def criar_recibo_fatura(request, id_cliente, id_fatura):
    if request.method != "POST":
        return JsonResponse({"error": "Método não permitido."}, status=405)

    metodo_nome = request.POST.get('metodo')
    if not metodo_nome:
        return JsonResponse({"error": "O método de pagamento é obrigatório."}, status=400)

    if not Modalidade.objects.filter(nome=metodo_nome).exists():
        return JsonResponse({"error": "Método de pagamento inválido."}, status=400)

    try:
        valor_str = request.POST.get('valor', '0').replace(',', '.')
        valor_pago = Decimal(valor_str)

        if valor_pago <= 0:
            return JsonResponse({"error": "O valor deve ser superior a zero."}, status=400)

        fatura = get_object_or_404(DocumentoFinalizado, id=id_fatura, cliente_id=id_cliente, empresa=request.empresa)

        nota_credito = DocumentoFinalizado.objects.filter(documento_origem=fatura, empresa=request.empresa).exclude(estado='Anulado').first()

        valor_nota_credito = nota_credito.valor_total if nota_credito else Decimal(0)

        divida_atual_com_desconto = fatura.valor_total - fatura.total_pago + valor_nota_credito

        if divida_atual_com_desconto <= 0:
            return JsonResponse({"error": "Esta fatura já se encontra liquidada."}, status=400)

        if valor_pago > divida_atual_com_desconto:
            return JsonResponse(
                {"error": f"O valor inserido ({valor_pago}) excede a dívida da fatura ({divida_atual_com_desconto})."}, status=400)

        ano_atual = timezone.now().year
        ultimo = Recibo.objects.filter(ano=ano_atual, empresa=request.empresa).order_by('-numero').first()
        novo_numero = (ultimo.numero + 1) if ultimo else 1

        # 2. Criar o Cabeçalho do Recibo
        recibo = Recibo.objects.create(
            tipo='RE',
            serie=str(ano_atual),
            numero=novo_numero,
            ano=ano_atual,
            cliente_id=id_cliente,
            data_emissao=timezone.now().date(),
            modalidade_nome=metodo_nome,
            valor_total=valor_pago,
            estado='Normal',
            empresa=request.empresa
        )

        # 3. Criar a Linha do Recibo para esta fatura única
        ReciboLinhas.objects.create(
            id_recibo_final=recibo,
            id_doc_final=fatura,
            documento_tipo=fatura.tipo,
            documento_numero=f"{fatura.serie}-{fatura.numero}/{fatura.ano}",
            data_emissao=fatura.data_emissao,
            valor_documento=fatura.valor_total,
            valor_recebido=valor_pago,
            valor_em_divida=divida_atual_com_desconto - valor_pago,
            empresa=request.empresa

        )

        # 4. Atualizar a Fatura
        fatura.total_pago += valor_pago
        fatura.save()

        return JsonResponse({
            "success": True,
            "message": f"Recibo emitido para a fatura {fatura.numero}.",
            "recibo_id": recibo.id
        })

    except IntegrityError:
        return JsonResponse({"error": "Erro de concorrência. Tente novamente."}, status=500)
    except Exception as e:
        return JsonResponse({"error": f"Erro interno: {str(e)}"}, status=500)


@login_required
@empresa_obrigatoria
def recibos_json(request):
    recibos = Recibo.objects.filter(empresa=request.empresa)

    ids_clientes = recibos.values_list('cliente_id', flat=True).distinct()

    clientes_map = dict(Cliente1.objects.filter(id_cliente__in=ids_clientes, empresa=request.empresa).values_list('id_cliente', 'nome'))

    data = []
    for r in recibos:
        data.append({
            "id_recibo": r.id,
            "documento": str(r),
            "data": r.data_emissao,
            "cliente": clientes_map.get(r.cliente_id, "Cliente não encontrado"),
            "total": float(r.valor_total),
            "estado": r.estado,
        })
    return JsonResponse(data, safe=False)

@login_required
@empresa_obrigatoria
def anular_recibo(request, id_recibo):
    if request.method != 'POST':
        return JsonResponse({'success': False, 'message': 'Método inválido.'}, status=405)

    motivo = request.POST.get('motivo', '').strip()
    if len(motivo) < 10 or len(motivo) > 255:
        return JsonResponse({'success': False, 'message': 'Insira o motivo com os parametros corretos.'}, status=400)

    try:
        with transaction.atomic():
            recibo = Recibo.objects.select_for_update().filter(id=id_recibo, empresa=request.empresa).first()

            if not recibo:
                return JsonResponse({'success': False, 'message': 'Recibo não encontrado.'}, status=404)

            if recibo.estado == 'anulado':
                return JsonResponse({'success': False, 'message': 'Este recibo já foi anulado anteriormente.'},
                                    status=400)

            # 2. Reverter os valores nos documentos finalizados
            linhas = ReciboLinhas.objects.filter(id_recibo_final=recibo, empresa=request.empresa)

            for linha in linhas:
                doc = linha.id_doc_final
                doc.total_pago -= linha.valor_recebido

                # O seu método save() no model já vai cuidar de:
                # - Se total_pago for 0 -> 'Pendente'
                # - Se total_pago < valor_total -> 'Parcial'
                doc.save()

            recibo.estado = 'Anulado'
            recibo.motivo_anulacao = motivo
            recibo.data_anulacao = timezone.now()
            recibo.save()

            return JsonResponse({
                'success': True,
                'message': f'Recibo {recibo.numero} anulado. Saldos dos documentos atualizados.'
            })

    except Exception as e:
        # Em caso de erro, o transaction.atomic faz o rollback de tudo
        return JsonResponse({'success': False, 'message': f'Erro ao processar: {str(e)}'}, status=500)


from django.shortcuts import get_object_or_404, redirect
from django.contrib import messages

@login_required
@empresa_obrigatoria
def emitir_nota_credito(request, fatura_id):
    fatura_original = get_object_or_404(DocumentoFinalizado, id=fatura_id, empresa=request.empresa)

    if fatura_original.tipo == 'NC':
        return JsonResponse({"erro": "Não é possível emitir uma Nota de Crédito de outra Nota de Crédito."}, status=400)

    if DocumentoFinalizado.objects.filter(documento_origem=fatura_original, tipo='NC', empresa=request.empresa).exclude(
            estado='Anulado').exists():
        return JsonResponse({"erro": "Já existe uma Nota de Crédito emitida para esta fatura."}, status=400)

    if DocumentoTemp.objects.filter(documento_origem=fatura_original, tipo='NC', empresa=request.empresa).exists():
        return JsonResponse({"erro": "Já existe uma Nota de Crédito em rascunho para esta fatura."}, status=400)

    recibos_ativos = ReciboLinhas.objects.filter(
        id_doc_final=fatura_original,
        id_recibo_final__estado='Normal',
        empresa=request.empresa
    )
    if recibos_ativos.exists():
        return JsonResponse({"erro": "Esta fatura possui recibos ativos. Deve anular os recibos antes de emitir uma Nota de Crédito."}, status=400)

    with transaction.atomic():
        ano = timezone.now().year
        contador, created = DocumentoContador.objects.select_for_update().get_or_create(
            tipo='NC',
            ano=ano,
            defaults={"serie": "NC", "ultimo_numero": 0},
            empresa = request.empresa
        )

        numero = contador.ultimo_numero + 1
        contador.ultimo_numero = numero
        contador.save()

        artigos_originais = FinArtigos.objects.filter(id_final=fatura_original, empresa=request.empresa)
        cliente_id = request.POST.get("cliente")
        try:
            cliente = Cliente1.objects.select_related(
                "transporte",
                "pagamento",
                "impostos"
            ).get(id_cliente=cliente_id,
                  empresa=request.empresa)
        except Cliente1.DoesNotExist:
            return JsonResponse({"erro": "Cliente inválido."}, status=400)

        try:
            nova_nc = DocumentoTemp.objects.create(
                tipo='NC',
                serie=contador.serie,
                numero=numero,
                ano=ano,
                cliente_id=fatura_original.cliente_id,
                valor_total=-fatura_original.valor_total,
                documento_origem=fatura_original,
                moeda_id=15,
                data_emissao=timezone.now(),
                data_vencimento=timezone.now(),
                transporte=cliente.transporte,
                impostos=cliente.impostos,
                estado='Rascunho',
                empresa=request.empresa
            )

            for art in artigos_originais:
                TempArtigos.objects.create(
                    id_temp=nova_nc,
                    id_art=art.id_art,
                    tipo=art.tipo,
                    descricao=art.descricao,
                    quantidade=art.quantidade,
                    preco=art.preco,
                    desconto=art.desconto,
                    taxa=art.taxa,
                    total=-abs(art.total),
                    motivo=art.motivo,
                    empresa = request.empresa
                )

            return JsonResponse({
                "id": nova_nc.id,
                "numero": nova_nc.numero,
                "cliente_id": nova_nc.cliente_id
            }, status=201)

        except Exception as e:
            return JsonResponse({"erro": str(e)}, status=400)

from django.db.models import Sum, Q
from django.db.models.functions import ExtractMonth, ExtractYear
from django.utils import timezone


@login_required
@empresa_obrigatoria
def dados_dashboard_ajax(request):
    hoje = timezone.now().date()
    # Pega o ano selecionado no select ou usa o ano atual por defeito
    param_ano = request.GET.get('ano')
    if param_ano and param_ano.isdigit():
        ano_selecionado = int(param_ano)
    else:
        ano_selecionado = hoje.year

    docs_validos = DocumentoFinalizado.objects.filter(
        empresa=request.empresa,
        estado='Finalizado',
    )

    # --- CÁLCULOS TOTAIS (Cards de cima) ---
    dados_faturado = docs_validos.aggregate(
        vendas=Sum('valor_total', filter=Q(tipo__in=['FT', 'FR', 'FS'])),
        devolucoes=Sum('valor_total', filter=Q(tipo='NC'))
    )
    total_faturado = (dados_faturado['vendas'] or 0) + (dados_faturado['devolucoes'] or 0)

    total_recibos = Recibo.objects.filter(empresa=request.empresa, estado='Normal').aggregate(total=Sum('valor_total'))[
                        'total'] or 0
    total_bruto_faturas = docs_validos.filter(tipo__in=['FT', 'FS', 'FR', 'NC']).aggregate(total=Sum('valor_total'))['total'] or 0
    saldo_em_divida = total_bruto_faturas - total_recibos

    total_vencidos_count = docs_validos.filter(
        tipo__in=['FT', 'FS', 'FR'],
        estado_pagamento__in=['Pendente', 'Parcial'],
        data_vencimento__lt=hoje
    ).count()

    # --- LÓGICA DO GRÁFICO (Agrupamento Mensal) ---
    # Inicializamos as listas com zeros para os 12 meses
    faturado_mes = [0] * 12
    pagos_mes = [0] * 12
    nao_pagos_mes = [0] * 12
    vencidos_mes = [0] * 12

    docs_ano = docs_validos.filter(
        data_vencimento__year=ano_selecionado,
        tipo__in=['FT', 'FR', 'FS', 'NC']
    ).annotate(mes=ExtractMonth('data_vencimento'))

    # 1. Primeiro, separa os documentos: Faturas e Notas de Crédito
    faturas = [d for d in docs_ano if d.tipo in ['FT', 'FR', 'FS']]
    notas_credito = [d for d in docs_ano if d.tipo == 'NC']

    # 2. Processa as Faturas e aplica as NCs
    for doc in faturas:
        idx = doc.mes - 1

        valor_total = float(doc.valor_total)
        valor_pago = float(doc.total_pago or 0)

        # Procura NCs associadas a esta fatura (assumindo que tens um campo documento_origem_id)
        # Se não tiveres este campo, terás de filtrar por cliente ou outra lógica
        valor_nc = sum([float(nc.valor_total) for nc in notas_credito if nc.documento_origem_id == doc.id])

        # O valor pendente real é: Fatura - Pago - (Valor da NC que é negativo)
        # Nota: como o valor_total da NC já é negativo (ex: -100), somamos o valor_total
        valor_pendente = valor_total - valor_pago + valor_nc

        faturado_mes[idx] += (valor_total + valor_nc)  # Faturado líquido
        pagos_mes[idx] += valor_pago

        if valor_pendente > 0:
            if doc.data_vencimento < hoje:
                vencidos_mes[idx] += valor_pendente
            else:
                nao_pagos_mes[idx] += valor_pendente

    anos_disponiveis = DocumentoFinalizado.objects.filter(
        empresa=request.empresa,
        tipo__in=['FT', 'FS', 'FR']
    ).annotate(ano_extraido=ExtractYear('data_vencimento')).values_list('ano_extraido', flat=True).distinct().order_by(
        '-ano_extraido')

    lista_anos = list(anos_disponiveis) if anos_disponiveis else [hoje.year]

    anos_recentes = list(range(ano_selecionado - 2, ano_selecionado + 1))

    dados_comparativos = {}

    for ano in anos_recentes:
        faturado_ano = [0] * 12
        docs_do_ano = docs_validos.filter(data_vencimento__year=ano, tipo__in=['FT', 'FR', 'FS','NC']).annotate(mes=ExtractMonth('data_vencimento'))

        for doc in docs_do_ano:
            idx = doc.mes - 1
            valor = float(doc.valor_total)
            faturado_ano[idx] += valor

        # Mantendo a estrutura exatamente como você precisa
        dados_comparativos[ano] = faturado_ano

    dados_totais_anuais = {}

    for ano in anos_recentes:
        # Filtramos os documentos deste ano específico
        docs_do_ano = docs_validos.filter(data_vencimento__year=ano)

        # Calculamos separadamente para evitar NULLs na subtração direta
        soma_vendas = docs_do_ano.filter(tipo__in=['FT', 'FR', 'FS']).aggregate(Sum('valor_total'))[
                          'valor_total__sum'] or 0
        soma_nc = docs_do_ano.filter(tipo='NC').aggregate(Sum('valor_total'))['valor_total__sum'] or 0

        dados_totais_anuais[str(ano)] = float(soma_vendas + soma_nc)


    anos_para_remover = [a for a in dados_totais_anuais.keys() if int(a) > ano_selecionado]
    for a in anos_para_remover:
        del dados_totais_anuais[a]

    top_clientes_ano = docs_validos.filter(
        data_vencimento__year=ano_selecionado,
        tipo__in=['FT', 'FR', 'FS', 'NC']
    ) \
        .values('cliente_nome') \
        .annotate(total_faturado=Sum('valor_total')) \
        .order_by('-total_faturado')[:5]  # Corrigido para 'total_faturado'

    nomes_top = [c['cliente_nome'] for c in top_clientes_ano]

    # 2. Histórico (total de todos os anos para esses clientes)
    historico_clientes = docs_validos.filter(
        tipo__in=['FT', 'FR', 'FS', 'NC'],
        cliente_nome__in=nomes_top
    ).values('cliente_nome') \
        .annotate(total_historico=Sum('valor_total'))

    mapa_historico = {c['cliente_nome']: float(c['total_historico'] or 0) for c in historico_clientes}

    # 3. Preparação para o JsonResponse
    labels_clientes = [c['cliente_nome'] for c in top_clientes_ano]
    valores_ano = [float(c['total_faturado'] or 0) for c in top_clientes_ano]
    valores_historico = [mapa_historico.get(nome, 0) for nome in labels_clientes]

    detalhes = docs_validos.filter(
        tipo__in=['FT', 'FR', 'FS', 'NC'],
    ).values(
        'id', 'tipo', 'numero', 'ano', 'cliente_nome', 'data_emissao', 'valor_total',
        'total_pago', 'estado_pagamento', 'data_vencimento', 'documento_origem_id'
    )

    empresa = request.empresa

    dados_empresa = {
        'id': empresa.id,
        'nome': empresa.nome,
        'nif': empresa.nif,
        'morada': empresa.morada,
        'codigo_postal': empresa.codigo_postal,
        'cidade': empresa.cidade,
        'pais': empresa.pais,
        'email': empresa.email,
        'telefone': empresa.telefone,
        'local': empresa.local,
    }
    return JsonResponse({
        "total_faturado": float(total_faturado),
        "saldo_pendente": float(max(0, saldo_em_divida)),
        "total_vencidos": total_vencidos_count,
        "ano_atual": ano_selecionado,
        "anos_lista": lista_anos,
        "chart_total": faturado_mes,
        "chart_pagos": pagos_mes,
        "chart_nao_pagos": nao_pagos_mes,
        "chart_vencidos": vencidos_mes,
        "mes_atual": hoje.month,
        "comparativo_anos": dados_comparativos,
        "totais_anuais": dados_totais_anuais,
        "top_clientes_labels": labels_clientes,
        "top_clientes_valores_ano": valores_ano,
        "top_clientes_valores_historico": valores_historico,
        "lista_detalhada": list(detalhes),
        'empresa': dados_empresa,
    })


@login_required
@empresa_obrigatoria
def editar_empresa_ajax(request, pk):
    try:
        empresa = Empresa.objects.get(pk=pk, user=request.user)
        data = json.loads(request.body)

        form = EmpresaForm(data, instance=empresa)

        if form.is_valid():
            form.save()  # Isto corre os métodos clean_nif e clean_codigo_postal
            return JsonResponse({'success': True, 'message': 'Guardado com sucesso!'})
        else:
            # Retorna os erros de validação para o JavaScript
            return JsonResponse({'success': False, 'error': form.errors.as_json()}, status=400)

    except Empresa.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Empresa não encontrada.'}, status=404)


@login_required
@empresa_obrigatoria
def adicionar_transporte_ajax(request):
    try:
        data = json.loads(request.body)
        descricao = data.get('descricao', '').strip().upper()

        if not descricao:
            return JsonResponse({'success': False, 'error': 'A matrícula é obrigatória.'}, status=400)

        clean_desc = descricao.replace('-', '').replace(' ', '')

        # O SAF-T não aceita caracteres especiais (@, #, !, etc.)
        if not re.match(r'^[A-Z0-9]{3,8}$', clean_desc):
            return JsonResponse({
                'success': False,
                'error': 'Formato de matrícula inválido para fins fiscais (SAFT).'
            }, status=400)

        if Transporte.objects.filter(empresa=request.empresa, descricao=descricao).exists():
            return JsonResponse({'success': False, 'error': 'Esta matrícula já está registada.'}, status=400)

        Transporte.objects.create(
            descricao=descricao,
            empresa=request.empresa
        )

        return JsonResponse({'success': True, 'message': 'Matrícula guardada com sucesso.'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)

@login_required
@empresa_obrigatoria
def obter_periodos_disponiveis(request):
    # Usamos campos diferentes para evitar o conflito
    periodos = DocumentoFinalizado.objects.filter(
        empresa=request.empresa
    ).exclude(
        tipo='GT'
    ).annotate(
        mes_doc=ExtractMonth('data_emissao'),
        ano_doc=ExtractYear('data_emissao')
    ).values('mes_doc', 'ano_doc').distinct().order_by('-ano_doc', '-mes_doc')

    # Ajuste: mapear o dicionário para a estrutura que o JS espera
    resultado = [
        {'mes': p['mes_doc'], 'ano': p['ano_doc']}
        for p in periodos
    ]

    return JsonResponse({'periodos': resultado})


from django.http import HttpResponse
import xml.etree.ElementTree as ET
from datetime import datetime


@login_required
@empresa_obrigatoria
def gerar_saft(request):
    mes = request.GET.get('mes')
    ano = request.GET.get('ano')

    docs = DocumentoFinalizado.objects.filter(
        empresa=request.empresa,
        data_emissao__month=mes,
        data_emissao__year=ano,
        estado='Finalizado'
    ).prefetch_related('artigos')

    # 1. Raiz e Header
    audit_file = ET.Element("AuditFile", xmlns="urn:OECD:StandardAuditFile-Tax:PT_3.01")
    header = ET.SubElement(audit_file, "Header")
    ET.SubElement(header, "AuditFileVersion").text = "3.01_01"
    ET.SubElement(header, "CompanyID").text = request.empresa.nif
    ET.SubElement(header, "TaxRegistrationNumber").text = request.empresa.nif
    ET.SubElement(header, "TaxAccountingBasis").text = "I"

    # 2. MasterFiles (Clientes únicos e Produtos únicos)
    master_files = ET.SubElement(audit_file, "MasterFiles")

    # Adicionar Clientes
    for cliente in docs.values('cliente_id', 'cliente_nome', 'cliente_contribuinte').distinct():
        cust = ET.SubElement(master_files, "Customer")
        ET.SubElement(cust, "CustomerID").text = str(cliente['cliente_id'])
        ET.SubElement(cust, "CustomerTaxID").text = cliente['cliente_contribuinte']
        ET.SubElement(cust, "CompanyName").text = cliente['cliente_nome']

    # 3. SourceDocuments
    source_docs = ET.SubElement(audit_file, "SourceDocuments")
    sales_invoices = ET.SubElement(source_docs, "SalesInvoices")

    for doc in docs:
        invoice = ET.SubElement(sales_invoices, "Invoice")
        ET.SubElement(invoice, "InvoiceNo").text = f"{doc.tipo} {doc.serie}/{doc.numero}"
        ET.SubElement(invoice, "InvoiceDate").text = str(doc.data_emissao)
        ET.SubElement(invoice, "CustomerID").text = str(doc.cliente_id)

        # Iterar sobre o teu modelo FinArtigos
        for linha in doc.artigos.all():
            line = ET.SubElement(invoice, "Line")
            ET.SubElement(line, "LineNumber").text = str(linha.id)
            ET.SubElement(line, "ProductCode").text = str(linha.id_art_id)
            ET.SubElement(line, "ProductDescription").text = linha.descricao
            ET.SubElement(line, "Quantity").text = str(linha.quantidade)
            ET.SubElement(line, "UnitPrice").text = str(linha.preco)
            ET.SubElement(line, "CreditAmount").text = str(linha.total)

        # Totais
        totals = ET.SubElement(invoice, "DocumentTotals")
        ET.SubElement(totals, "GrossTotal").text = str(doc.valor_total)
        ET.SubElement(totals, "NetTotal").text = str(doc.valor_total)  # Simplificação: considera IVA à parte

    payments = ET.SubElement(source_docs, "Payments")

    # Buscar recibos do período
    recibos = Recibo.objects.filter(
        empresa=request.empresa,
        data_emissao__month=mes,
        data_emissao__year=ano,
        estado='Normal'
    ).prefetch_related('linhas')

    for recibo in recibos:
        payment = ET.SubElement(payments, "Payment")
        ET.SubElement(payment, "PaymentRefNo").text = f"{recibo.tipo} {recibo.serie}/{recibo.numero}"
        ET.SubElement(payment, "Period").text = str(recibo.data_emissao.month)
        ET.SubElement(payment, "TransactionDate").text = str(recibo.data_emissao)
        ET.SubElement(payment, "PaymentType").text = "RC"  # RC para recibo
        ET.SubElement(payment, "Description").text = recibo.modalidade_nome or "Pagamento"
        ET.SubElement(payment, "SystemEntryDate").text = recibo.criado_em.strftime('%Y-%m-%dT%H:%M:%S')
        ET.SubElement(payment, "CustomerID").text = str(recibo.cliente_id)

        # Documentos liquidados (Crucial para a AT!)
        for linha in recibo.linhas.all():
            doc_ref = ET.SubElement(payment, "DocumentTotals")
            # O SAF-T exige saber qual a fatura que foi paga
            source_doc = ET.SubElement(doc_ref, "SourceDocumentID")
            ET.SubElement(source_doc, "OriginatingON").text = f"{linha.documento_tipo} {linha.documento_numero}"

            ET.SubElement(doc_ref, "AmountReceived").text = str(linha.valor_recebido)

        # Total do Recibo
        payment_totals = ET.SubElement(payment, "PaymentMethod")
        ET.SubElement(payment_totals, "PaymentMethod").text = "OU"  # Outros (ex: Transferência)
        ET.SubElement(payment_totals, "PaymentAmount").text = str(recibo.valor_total)
    # 4. Finalização
    tree = ET.ElementTree(audit_file)
    response = HttpResponse(content_type="application/xml")
    response['Content-Disposition'] = f'attachment; filename="SAFT_{mes}_{ano}.xml"'
    tree.write(response, encoding='utf-8', xml_declaration=True)
    return response